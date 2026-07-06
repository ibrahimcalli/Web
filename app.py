"""
Portföy Gayrimenkul Web Sistemi - Backend
FastAPI + SQLite | JWT Auth | Belge Parser
"""

from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form, status, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from jose import JWTError, jwt
from passlib.context import CryptContext
from datetime import datetime, timedelta
from typing import Optional, List, TYPE_CHECKING
from pydantic import BaseModel
import sqlite3, json, os, shutil, uuid, re, time, secrets, hashlib
from collections import defaultdict
import requests as http_requests
try:
    from PIL import Image
    PIL_VAR = True
except ImportError:
    PIL_VAR = False
from pathlib import Path

# ─── Konfig ────────────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent
DB_PATH    = BASE_DIR / "emlak_web.db"
UPLOAD_DIR = BASE_DIR / "static" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ── Güvenlik ── .env dosyasından oku, yoksa varsayılan (prod'da mutlaka değiştirin)
SECRET_KEY   = os.environ.get("SECRET_KEY", "emlak-gizli-anahtar-2026-degistir")
ALGORITHM    = "HS256"
TOKEN_EXPIRE = 60 * 24  # dakika

# ── Rate Limiting (bellek tabanlı, sunucu restart'ta sıfırlanır) ──────────────
_giris_denemeleri: dict = defaultdict(list)   # ip → [timestamp, ...]
_engellenen_ipler: dict = {}                   # ip → engel_bitis_zamani
MAX_DENEME      = 5     # Bu kadar başarısız denemeden sonra...
ENGEL_SURESI    = 900   # ...bu kadar saniye (15 dk) engelle
PENCERE_SURESI  = 300   # 5 dakika pencere içinde MAX_DENEME aşılırsa engelle

# ─── Kategori Ağacı (masaüstünden alındı) ─────────────────────────────────────
KATEGORILER = {
    "Konut":           ["Satılık", "Kiralık", "Devren Satılık", "Devren Kiralık"],
    "İş Yeri":         ["Satılık", "Kiralık", "Devren Satılık", "Devren Kiralık"],
    "Arsa":            ["Satılık", "Kiralık"],
    "Konut Projeleri": ["Satılık"],
    "Bina":            ["Satılık", "Kiralık"],
    "Devre Mülk":      ["Satılık", "Kiralık"],
    "Turistik Tesis":  ["Satılık", "Kiralık"],
}

ILAN_TIPLERI = {
    "Konut": {
        "Satılık":        ["Daire", "Rezidans", "Villa", "Yazlık", "Müstakil Ev", "Kooperatif", "Bungalov"],
        "Kiralık":        ["Daire", "Rezidans", "Villa", "Yazlık", "Müstakil Ev"],
        "Devren Satılık": ["Daire", "Rezidans", "Villa", "Müstakil Ev"],
        "Devren Kiralık": ["Daire", "Rezidans", "Villa", "Müstakil Ev"],
    },
    "İş Yeri": {
        "Satılık":        ["Ofis", "Dükkan/Mağaza", "Depo/Antrepo", "Atölye", "Fabrika", "Akaryakıt İstasyonu"],
        "Kiralık":        ["Ofis", "Dükkan/Mağaza", "Depo/Antrepo", "Atölye", "Fabrika"],
        "Devren Satılık": ["Ofis", "Dükkan/Mağaza", "Atölye"],
        "Devren Kiralık": ["Ofis", "Dükkan/Mağaza", "Atölye"],
    },
    "Arsa": {
        "Satılık": ["Konut Arsası", "Ticari Arsa", "Tarla", "Bahçe", "Bağ", "Çiftlik"],
        "Kiralık": ["Tarla", "Bahçe", "Bağ"],
    },
    "Konut Projeleri": {
        "Satılık": ["Daire", "Villa", "Rezidans", "Müstakil"],
    },
    "Bina": {
        "Satılık": ["Apartman", "İş Merkezi", "Fabrika", "Otel"],
        "Kiralık": ["Apartman", "İş Merkezi"],
    },
    "Devre Mülk": {
        "Satılık": ["Devre Mülk"],
        "Kiralık": ["Devre Mülk"],
    },
    "Turistik Tesis": {
        "Satılık": ["Otel", "Pansiyon", "Tatil Köyü", "Kamp Alanı"],
        "Kiralık": ["Otel", "Pansiyon", "Tatil Köyü", "Kamp Alanı"],
    },
}

# ─── Form Alan Şablonları ──────────────────────────────────────────────────────
ALAN_SABLONLARI = {
    "konut_satilik": [
        {"key": "net_m2",        "label": "Net m²",          "type": "number"},
        {"key": "brut_m2",       "label": "Brüt m²",         "type": "number"},
        {"key": "oda_sayisi",    "label": "Oda Sayısı",       "type": "select",
         "options": ["1+0","1+1","2+0","2+1","3+1","3+2","4+1","4+2","5+1","5+2","6+"]},
        {"key": "bina_kati",     "label": "Bina Katı",        "type": "number"},
        {"key": "bulundugu_kat", "label": "Bulunduğu Kat",    "type": "number"},
        {"key": "bina_yasi",     "label": "Bina Yaşı",        "type": "number"},
        {"key": "isitma",        "label": "Isıtma",           "type": "select",
         "options": ["Doğalgaz (Kombi)","Doğalgaz (Merkezi)","Klima","Elektrikli","Soba","Yerden Isıtma","Yok"]},
        {"key": "banyo_sayisi",  "label": "Banyo Sayısı",     "type": "number"},
        {"key": "tapu_durumu",   "label": "Tapu Durumu",      "type": "select",
         "options": ["Kat Mülkiyeti","Kat İrtifakı","Arsa Tapusu","Hisseli Tapu"]},
        {"key": "krediye_uygun", "label": "Krediye Uygun",    "type": "select",
         "options": ["Var","Yok"]},
        {"key": "cephe",         "label": "Cephe",            "type": "text"},
        {"key": "esyali",        "label": "Eşyalı",           "type": "select",
         "options": ["Yok","Var","Yarı Eşyalı"]},
        {"key": "balkon",        "label": "Balkon",           "type": "select",
         "options": ["Var","Yok"]},
        {"key": "asansor",       "label": "Asansör",          "type": "select",
         "options": ["Var","Yok"]},
        {"key": "otopark",       "label": "Otopark",          "type": "select",
         "options": ["Var","Yok","Açık","Kapalı"]},
        {"key": "site_icinde",   "label": "Site İçinde",      "type": "select",
         "options": ["Evet","Hayır"]},
        {"key": "kullanim",      "label": "Kullanım Durumu",  "type": "select",
         "options": ["Boş","Kiracılı","Mal Sahibi"]},
        {"key": "takas",         "label": "Takas",            "type": "select",
         "options": ["Var","Yok"]},
        {"key": "ada",           "label": "Ada",              "type": "text"},
        {"key": "parsel",        "label": "Parsel",           "type": "text"},
        {"key": "ozellikler",    "label": "Değer Katan Özellikler", "type": "textarea"},
    ],
    "konut_kiralik": [
        {"key": "net_m2",        "label": "Net m²",           "type": "number"},
        {"key": "brut_m2",       "label": "Brüt m²",          "type": "number"},
        {"key": "oda_sayisi",    "label": "Oda Sayısı",        "type": "select",
         "options": ["1+0","1+1","2+0","2+1","3+1","3+2","4+1","4+2","5+1","5+2","6+"]},
        {"key": "bina_kati",     "label": "Bina Katı",         "type": "number"},
        {"key": "bulundugu_kat", "label": "Bulunduğu Kat",     "type": "number"},
        {"key": "bina_yasi",     "label": "Bina Yaşı",         "type": "number"},
        {"key": "isitma",        "label": "Isıtma",            "type": "select",
         "options": ["Doğalgaz (Kombi)","Doğalgaz (Merkezi)","Klima","Elektrikli","Soba","Yerden Isıtma","Yok"]},
        {"key": "banyo_sayisi",  "label": "Banyo Sayısı",      "type": "number"},
        {"key": "esyali",        "label": "Eşyalı",            "type": "select",
         "options": ["Yok","Var","Yarı Eşyalı"]},
        {"key": "balkon",        "label": "Balkon",            "type": "select",
         "options": ["Var","Yok"]},
        {"key": "asansor",       "label": "Asansör",           "type": "select",
         "options": ["Var","Yok"]},
        {"key": "otopark",       "label": "Otopark",           "type": "select",
         "options": ["Var","Yok"]},
        {"key": "site_icinde",   "label": "Site İçinde",       "type": "select",
         "options": ["Evet","Hayır"]},
        {"key": "depozito",      "label": "Depozito (ay)",     "type": "number"},
        {"key": "aidat",         "label": "Aidat (TL)",        "type": "number"},
        {"key": "ozellikler",    "label": "Değer Katan Özellikler","type": "textarea"},
    ],
    "isyeri_satilik": [
        {"key": "net_m2",        "label": "Net m²",            "type": "number"},
        {"key": "brut_m2",       "label": "Brüt m²",           "type": "number"},
        {"key": "kat",           "label": "Kat",               "type": "number"},
        {"key": "bina_yasi",     "label": "Bina Yaşı",         "type": "number"},
        {"key": "isitma",        "label": "Isıtma",            "type": "select",
         "options": ["Klima","Doğalgaz (Kombi)","Doğalgaz (Merkezi)","Elektrikli","Soba","Yok"]},
        {"key": "tapu_durumu",   "label": "Tapu Durumu",       "type": "select",
         "options": ["Kat Mülkiyeti","Kat İrtifakı","Arsa Tapusu","Hisseli Tapu"]},
        {"key": "krediye_uygun", "label": "Krediye Uygun",     "type": "select",
         "options": ["Var","Yok"]},
        {"key": "kullanim",      "label": "Kullanım Durumu",   "type": "select",
         "options": ["Boş","Kiracılı","Mal Sahibi"]},
        {"key": "takas",         "label": "Takas",             "type": "select",
         "options": ["Var","Yok"]},
        {"key": "cephe",         "label": "Cephe",             "type": "text"},
        {"key": "asansor",       "label": "Asansör",           "type": "select",
         "options": ["Var","Yok"]},
        {"key": "otopark",       "label": "Otopark",           "type": "select",
         "options": ["Var","Yok","Açık","Kapalı"]},
        {"key": "ozellikler",    "label": "Özellikler",        "type": "textarea"},
    ],
    "arsa": [
        {"key": "alan_m2",       "label": "Alan (m²)",         "type": "number"},
        {"key": "ada",           "label": "Ada",               "type": "text"},
        {"key": "parsel",        "label": "Parsel",            "type": "text"},
        {"key": "kaks",          "label": "KAKS/EMSAL",        "type": "text"},
        {"key": "taks",          "label": "TAKS",              "type": "text"},
        {"key": "imar_durumu",   "label": "İmar Durumu",       "type": "select",
         "options": ["Konut İmarlı","Ticari İmarlı","Tarım","Orman","İmarsız","Plansız"]},
        {"key": "tapu_durumu",   "label": "Tapu Durumu",       "type": "select",
         "options": ["Arsa Tapusu","Hisseli Tapu","Tarla Tapusu"]},
        {"key": "takas",         "label": "Takas",             "type": "select",
         "options": ["Var","Yok"]},
        {"key": "ozellikler",    "label": "Özellikler",        "type": "textarea"},
    ],
    "turistik": [
        {"key": "net_m2",        "label": "Net m²",            "type": "number"},
        {"key": "oda_sayisi",    "label": "Oda/Suit Sayısı",   "type": "number"},
        {"key": "yatak_kapasitesi","label":"Yatak Kapasitesi", "type": "number"},
        {"key": "yildiz",        "label": "Yıldız",            "type": "select",
         "options": ["1","2","3","4","5","Butik","Belspaş","Pansiyon"]},
        {"key": "havuz",         "label": "Havuz",             "type": "select",
         "options": ["Var","Yok","Kapalı","Açık"]},
        {"key": "plaj",          "label": "Plaj/Deniz",        "type": "select",
         "options": ["Denize Sıfır","Yakın","Uzak"]},
        {"key": "tapu_durumu",   "label": "Tapu Durumu",       "type": "select",
         "options": ["Kat Mülkiyeti","Arsa Tapusu","Hisseli Tapu"]},
        {"key": "ozellikler",    "label": "Özellikler",        "type": "textarea"},
    ],
    "genel": [
        {"key": "net_m2",        "label": "Net m²",            "type": "number"},
        {"key": "brut_m2",       "label": "Brüt m²",           "type": "number"},
        {"key": "tapu_durumu",   "label": "Tapu Durumu",       "type": "select",
         "options": ["Kat Mülkiyeti","Arsa Tapusu","Hisseli Tapu"]},
        {"key": "ozellikler",    "label": "Özellikler",        "type": "textarea"},
    ],
}

def alan_sablonu_sec(ana_kat: str, alt_kat: str, ilan_tipi: str = "") -> list:
    """Kategori + alt kategoriye göre uygun alan şablonu seç."""
    kiralik = alt_kat in ("Kiralık", "Devren Kiralık")
    if "Arsa" in ana_kat:
        return ALAN_SABLONLARI["arsa"]
    elif "Turistik" in ana_kat:
        return ALAN_SABLONLARI["turistik"]
    elif "İş" in ana_kat or "Ticari" in ana_kat:
        return ALAN_SABLONLARI["isyeri_satilik"]
    elif "Konut" in ana_kat:
        return ALAN_SABLONLARI["konut_kiralik"] if kiralik else ALAN_SABLONLARI["konut_satilik"]
    return ALAN_SABLONLARI["genel"]

# ─── Auth ──────────────────────────────────────────────────────────────────────
pwd_ctx = CryptContext(schemes=["bcrypt"], deprecated="auto")

def istek_ip_al(request=None):
    """İstek IP adresini güvenli şekilde al."""
    from fastapi import Request
    return "bilinmiyor"

def rate_limit_kontrol(ip: str) -> None:
    """Brute force saldırısı var mı kontrol et. Varsa 429 fırlat."""
    simdi = time.time()

    # Engel süresi bitti mi?
    if ip in _engellenen_ipler:
        if simdi < _engellenen_ipler[ip]:
            kalan = int(_engellenen_ipler[ip] - simdi)
            raise HTTPException(
                status_code=429,
                detail=f"Çok fazla başarısız deneme. {kalan} saniye bekleyin.",
                headers={"Retry-After": str(kalan)}
            )
        else:
            del _engellenen_ipler[ip]
            _giris_denemeleri[ip] = []

    # Pencere içindeki denemeleri temizle
    pencere_baslangic = simdi - PENCERE_SURESI
    _giris_denemeleri[ip] = [t for t in _giris_denemeleri[ip] if t > pencere_baslangic]

def rate_limit_basarisiz(ip: str) -> None:
    """Başarısız giriş kaydı. Limit aşılırsa engelle."""
    simdi = time.time()
    _giris_denemeleri[ip].append(simdi)
    if len(_giris_denemeleri[ip]) >= MAX_DENEME:
        _engellenen_ipler[ip] = simdi + ENGEL_SURESI
        raise HTTPException(
            status_code=429,
            detail=f"Çok fazla başarısız deneme. 15 dakika bekleyin."
        )

def rate_limit_basarili(ip: str) -> None:
    """Başarılı girişte sayacı sıfırla."""
    _giris_denemeleri.pop(ip, None)
    _engellenen_ipler.pop(ip, None)

oauth2  = OAuth2PasswordBearer(tokenUrl="/api/auth/giris", auto_error=False)

def hash_sifre(sifre): return pwd_ctx.hash(sifre)
def sifre_dogrula(plain, hashed): return pwd_ctx.verify(plain, hashed)

def token_olustur(data: dict, dakika: int = TOKEN_EXPIRE):
    exp = datetime.utcnow() + timedelta(minutes=dakika)
    return jwt.encode({**data, "exp": exp}, SECRET_KEY, algorithm=ALGORITHM)

def token_coz(token: str = Depends(oauth2)):
    if not token:
        return None
    try:
        return jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
    except JWTError:
        return None

def admin_gerek(payload=Depends(token_coz)):
    if not payload or payload.get("rol") != "admin":
        raise HTTPException(status_code=401, detail="Yetkisiz erişim")
    return payload

def kullanici_gerek(payload=Depends(token_coz)):
    if not payload:
        raise HTTPException(status_code=401, detail="Giriş yapmanız gerekiyor")
    # Admin değilse onay kontrolü
    if payload.get("rol") != "admin":
        conn = get_db()
        k = conn.execute("SELECT onay, aktif FROM kullanicilar WHERE email=?", (payload["sub"],)).fetchone()
        conn.close()
        if not k or not k["aktif"]:
            raise HTTPException(status_code=403, detail="Hesabınız devre dışı")
        if not k["onay"]:
            raise HTTPException(status_code=403, detail="Hesabınız henüz onaylanmadı. Lütfen yönetici onayını bekleyin.")
    return payload

# ─── Veritabanı ────────────────────────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()

    c.executescript("""
    CREATE TABLE IF NOT EXISTS kullanicilar (
        id            INTEGER PRIMARY KEY AUTOINCREMENT,
        ad_soyad      TEXT NOT NULL,
        email         TEXT UNIQUE NOT NULL,
        sifre         TEXT NOT NULL,
        rol           TEXT DEFAULT 'kullanici',
        aktif         INTEGER DEFAULT 1,
        onay_durumu   TEXT DEFAULT 'bekliyor',
        profil_resmi  TEXT DEFAULT '',
        olusturma     TEXT DEFAULT (datetime('now'))
    );
    -- Migration: profil_resmi kolonu yoksa ekle
    CREATE TABLE IF NOT EXISTS _migrasyon_yapildi (ad TEXT PRIMARY KEY);

    CREATE TABLE IF NOT EXISTS bannerlar (
        id          INTEGER PRIMARY KEY AUTOINCREMENT,
        baslik      TEXT DEFAULT '',
        alt_metin   TEXT DEFAULT '',
        link_url    TEXT DEFAULT '',
        link_hedef  TEXT DEFAULT '_self',
        resim_url   TEXT DEFAULT '',
        tip         TEXT DEFAULT 'slider',
        konum       TEXT DEFAULT 'anasayfa_hero_alti',
        boyut       TEXT DEFAULT 'genis',
        sira        INTEGER DEFAULT 0,
        aktif       INTEGER DEFAULT 1,
        olusturma   TEXT DEFAULT (datetime('now'))
    );
    

    CREATE TABLE IF NOT EXISTS portfoyler (
        id              INTEGER PRIMARY KEY AUTOINCREMENT,
        baslik          TEXT NOT NULL,
        ana_kategori    TEXT NOT NULL,
        alt_kategori    TEXT NOT NULL,
        ilan_tipi       TEXT,
        il              TEXT DEFAULT 'Mugla',
        ilce            TEXT DEFAULT 'Fethiye',
        mahalle         TEXT,
        fiyat           TEXT,
        para_birimi     TEXT DEFAULT 'TL',
        aciklama        TEXT,
        saha_notu       TEXT,
        gps             TEXT,
        durum           TEXT DEFAULT 'Taslak',
        alanlar         TEXT DEFAULT '{}',
        resimler        TEXT DEFAULT '[]',
        musteri_ad      TEXT,
        musteri_tel     TEXT,
        musteri_mail    TEXT,
        musteri_adres   TEXT DEFAULT '',
        musteri_tc      TEXT DEFAULT '',
        musteri_not     TEXT DEFAULT '',
        sahip_goster    INTEGER DEFAULT 0,
        kaynak          TEXT DEFAULT 'web',
        olusturma       TEXT DEFAULT (datetime('now')),
        guncelleme      TEXT DEFAULT (datetime('now'))
    );

    CREATE TABLE IF NOT EXISTS kullanici_istekleri (
        id           INTEGER PRIMARY KEY AUTOINCREMENT,
        ad_soyad     TEXT,
        telefon      TEXT,
        email        TEXT,
        mesaj        TEXT,
        portfoy_id   INTEGER,
        durum        TEXT DEFAULT 'Yeni',
        olusturma    TEXT DEFAULT (datetime('now')),
        FOREIGN KEY(portfoy_id) REFERENCES portfoyler(id)
    );

    CREATE TABLE IF NOT EXISTS site_ayarlari (
        anahtar TEXT PRIMARY KEY,
        deger   TEXT
    );
    CREATE TABLE IF NOT EXISTS bannerlar (
        id          INTEGER PRIMARY KEY AUTOINCREMENT,
        tip         TEXT DEFAULT 'slider',  -- slider | duyuru
        baslik      TEXT DEFAULT '',
        aciklama    TEXT DEFAULT '',
        resim_url   TEXT DEFAULT '',
        link_url    TEXT DEFAULT '',
        link_metin  TEXT DEFAULT '',
        renk_arka   TEXT DEFAULT '',
        renk_metin  TEXT DEFAULT '#ffffff',
        konum       TEXT DEFAULT 'ana_hero_alti', -- ana_hero_alti | tum_nav_alti | ana_ust | ilan_arasi
        boyut       TEXT DEFAULT 'genis',   -- tam | genis | orta | kucuk
        sira        INTEGER DEFAULT 0,
        aktif       INTEGER DEFAULT 1,
        olusturma   TEXT DEFAULT (datetime('now'))
    );

    CREATE TABLE IF NOT EXISTS blog_yazilari (
        id          INTEGER PRIMARY KEY AUTOINCREMENT,
        baslik      TEXT NOT NULL,
        slug        TEXT UNIQUE,
        icerik      TEXT DEFAULT '',
        ozet        TEXT DEFAULT '',
        etiketler   TEXT DEFAULT '[]',
        kapak_resim TEXT DEFAULT '',
        durum       TEXT DEFAULT 'Taslak',
        yazar_id    INTEGER,
        olusturma   TEXT DEFAULT (datetime('now')),
        guncelleme  TEXT DEFAULT (datetime('now'))
    );
    """)

    # Varsayılan admin
    row = c.execute("SELECT id FROM kullanicilar WHERE email='bilgi@portfoygayrimenkul.com.tr'").fetchone()
    if not row:
        c.execute("""INSERT INTO kullanicilar (ad_soyad,email,sifre,rol)
                     VALUES (?,?,?,?)""",
                  ("Portföy Gayrimenkul", "bilgi@portfoygayrimenkul.com.tr",
                   hash_sifre("admin123"), "admin"))

    # Varsayılan site ayarları
    defaults = {
        "site_adi":    "Portföy Gayrimenkul",
        "site_slogan": "Fethiye'nin Güvenilir Gayrimenkul Danışmanı",
        "telefon":     "0542 966 36 36",
        "email":       "bilgi@portfoygayrimenkul.com.tr",
        "adres":       "Fethiye / Muğla",
        "web_sitesi":  "portfoygayrimenkul.com.tr",
        "renk_tema":   "",
        "logo_url":    "",
        "sosyal_ig":   "",
        "sosyal_fb":   "",
        "sosyal_wa":   "",
    }
    for k, v in defaults.items():
        c.execute("INSERT OR IGNORE INTO site_ayarlari VALUES (?,?)", (k, v))

    # Migration
    for mig in [
        "ALTER TABLE kullanicilar ADD COLUMN profil_resmi TEXT DEFAULT ''",
        "ALTER TABLE kullanicilar ADD COLUMN onay INTEGER DEFAULT 1",
        "ALTER TABLE portfoyler ADD COLUMN sahip_goster INTEGER DEFAULT 0",
        "ALTER TABLE portfoyler ADD COLUMN musteri_tc TEXT DEFAULT ''",
        "ALTER TABLE portfoyler ADD COLUMN musteri_adres TEXT DEFAULT ''",
        "ALTER TABLE portfoyler ADD COLUMN musteri_notlar TEXT DEFAULT ''",
        "ALTER TABLE portfoyler ADD COLUMN musteri_adres TEXT DEFAULT ''",
        "ALTER TABLE portfoyler ADD COLUMN musteri_tc TEXT DEFAULT ''",
        "ALTER TABLE portfoyler ADD COLUMN musteri_not TEXT DEFAULT ''",
        """CREATE TABLE IF NOT EXISTS bannerlar (id INTEGER PRIMARY KEY AUTOINCREMENT, tip TEXT DEFAULT 'slider', baslik TEXT DEFAULT '', aciklama TEXT DEFAULT '', resim_url TEXT DEFAULT '', link_url TEXT DEFAULT '', link_metin TEXT DEFAULT '', renk_arka TEXT DEFAULT '', renk_metin TEXT DEFAULT '#ffffff', konum TEXT DEFAULT 'ana_hero_alti', boyut TEXT DEFAULT 'genis', sira INTEGER DEFAULT 0, aktif INTEGER DEFAULT 1, olusturma TEXT DEFAULT (datetime('now')))""",
    ]:
        try:
            c.execute(mig)
        except Exception:
            pass

    conn.commit()
    conn.close()

# ─── Belge Parser ──────────────────────────────────────────────────────────────
def docx_parse(dosya_yolu: str) -> dict:
    """docx dosyasından portföy alanlarını çıkar."""
    try:
        from docx import Document
        doc = Document(dosya_yolu)
        tablo_verisi = {}
        metin_bloklari = []

        # Tabloları oku
        for tablo in doc.tables:
            for satir in tablo.rows:
                cells = [c.text.strip() for c in satir.cells]
                if len(cells) >= 2 and cells[0] and cells[1]:
                    tablo_verisi[cells[0].strip(":")] = cells[1]

        # Paragrafları oku
        for para in doc.paragraphs:
            metin = para.text.strip()
            if metin:
                metin_bloklari.append(metin)

        return _tablo_to_portfoy(tablo_verisi, metin_bloklari)
    except Exception as e:
        return {"hata": str(e)}

def html_parse(icerik: str) -> dict:
    """HTML içeriğinden portföy alanlarını çıkar."""
    from html.parser import HTMLParser
    import html as html_module

    tablo_verisi = {}
    metin_bloklari = []

    # Tablo hücrelerini yakala
    td_pattern = re.findall(r'<td[^>]*>(.*?)</td>', icerik, re.DOTALL | re.IGNORECASE)
    temiz = [re.sub(r'<[^>]+>', '', h).strip() for h in td_pattern]
    temiz = [html_module.unescape(t) for t in temiz if t]

    for i in range(0, len(temiz) - 1, 2):
        if temiz[i] and temiz[i+1]:
            tablo_verisi[temiz[i].strip(":")] = temiz[i+1]

    # h2 başlıklar (bölüm bilgisi)
    for m in re.finditer(r'<h[1-3][^>]*>(.*?)</h[1-3]>', icerik, re.DOTALL | re.IGNORECASE):
        t = re.sub(r'<[^>]+>', '', m.group(1)).strip()
        if t:
            metin_bloklari.append(t)

    return _tablo_to_portfoy(tablo_verisi, metin_bloklari)

def _tablo_to_portfoy(tablo: dict, metinler: list) -> dict:
    """Tablo verisi + metin → portföy dict."""
    def bul(*anahtarlar):
        for a in anahtarlar:
            for k, v in tablo.items():
                if a.lower() in k.lower():
                    return v.strip()
        return ""

    # Başlık: ilk metin satırı genellikle işletme adı + ilan tipi
    baslik = ""
    for m in metinler:
        if len(m) > 5 and not m.startswith("IC "):
            baslik = m
            break

    # Kategori tespiti
    kat_str  = bul("Kategori", "Kategori")
    alt_str  = bul("Ilan Tipi", "İlan Tipi", "Alt Kategori")
    durum_str = bul("Durum", "Statü")

    # Kategori → ana_kat / alt_kat çıkar
    ana_kat, alt_kat, ilan_tipi = "Konut", "Satılık", ""
    for k in KATEGORILER:
        if k.lower() in kat_str.lower():
            ana_kat = k
            break
    for a in KATEGORILER.get(ana_kat, []):
        if a.lower() in kat_str.lower() or a.lower() in alt_str.lower():
            alt_kat = a
            break

    # Fiyat: ilk büyük sayıyı yakala
    fiyat = ""
    for m in metinler:
        match = re.search(r'[\d.,]+\s*(TL|EUR|USD|\$|€)?', m)
        if match and any(c.isdigit() for c in match.group()):
            sayi = re.sub(r'[.,\s]', '', match.group())
            if len(sayi) >= 4:
                fiyat = match.group().strip()
                break

    # GPS
    gps = bul("GPS", "Konum", "google")
    gps_match = re.search(r'q=([\d.]+),([\d.]+)', gps)
    if gps_match:
        gps = f"{gps_match.group(1)},{gps_match.group(2)}"

    # Açıklama + saha notu
    aciklama = ""
    saha_notu = ""
    yakalanmis = False
    for m in metinler:
        if "açıklama" in m.lower() or "ACIKLAMA" in m:
            yakalanmis = True
            continue
        if "saha" in m.lower() or "SAHA" in m:
            yakalanmis = False
            saha_notu = m
            continue
        if yakalanmis and not aciklama:
            aciklama = m

    # Dinamik alanlar
    alanlar = {}
    alan_map = {
        "net_m2":        ["Net M2", "Net m2", "Net M²"],
        "brut_m2":       ["Brut M2", "Brüt m2", "Brüt M²"],
        "oda_sayisi":    ["Oda Sayisi", "Oda Sayısı"],
        "bina_kati":     ["Bina Kati", "Bina Katı"],
        "bulundugu_kat": ["Kat", "Bulunduğu Kat"],
        "bina_yasi":     ["Bina Yasi", "Bina Yaşı"],
        "banyo_sayisi":  ["Banyo Sayisi", "Banyo Sayısı"],
        "isitma":        ["Isitma", "Isıtma"],
        "tapu_durumu":   ["Tapu Durumu"],
        "cephe":         ["Cephe"],
        "esyali":        ["Esyali", "Eşyalı"],
        "balkon":        ["Balkon"],
        "asansor":       ["Asansor", "Asansör"],
        "otopark":       ["Otopark"],
        "site_icinde":   ["Site Icinde", "Site İçinde"],
        "kullanim":      ["Kullanim", "Kullanım"],
        "takas":         ["Takas"],
        "ada":           ["Ada"],
        "parsel":        ["Parsel"],
        "krediye_uygun": ["Krediye Uygun"],
        "kimden":        ["Kimden"],
        "ozellikler":    ["Deger Katan", "Değer Katan", "Özellikler"],
    }
    for hedef, kaynaklar in alan_map.items():
        deger = bul(*kaynaklar)
        if deger:
            alanlar[hedef] = deger

    musteri_ad   = bul("Ad Soyad", "Sahip", "Müşteri")
    musteri_tel  = bul("Telefon", "Tel")
    musteri_mail = bul("E-posta", "Mail", "Email")
    il_ilce      = bul("Il / Ilce", "İl / İlçe", "İl/İlçe")
    il, ilce     = "Muğla", "Fethiye"
    if "/" in il_ilce:
        parcalar = [p.strip() for p in il_ilce.split("/")]
        il  = parcalar[0] if parcalar[0] else il
        ilce = parcalar[1] if len(parcalar) > 1 and parcalar[1] else ilce
    mahalle = bul("Mahalle", "Bölge", "Semt")

    return {
        "baslik":       baslik,
        "ana_kategori": ana_kat,
        "alt_kategori": alt_kat,
        "ilan_tipi":    ilan_tipi,
        "fiyat":        fiyat,
        "il":           il,
        "ilce":         ilce,
        "mahalle":      mahalle,
        "gps":          gps,
        "aciklama":     aciklama,
        "saha_notu":    saha_notu,
        "musteri_ad":   musteri_ad,
        "musteri_tel":  musteri_tel,
        "musteri_mail": musteri_mail,
        "alanlar":      alanlar,
        "kaynak":       "belge_import",
    }

# ─── Pydantic Modeller ─────────────────────────────────────────────────────────
class PortfoyGiren(BaseModel):
    baslik:       str
    ana_kategori: str
    alt_kategori: str
    ilan_tipi:    Optional[str] = ""
    il:           Optional[str] = "Muğla"
    ilce:         Optional[str] = "Fethiye"
    mahalle:      Optional[str] = ""
    fiyat:        Optional[str] = ""
    para_birimi:  Optional[str] = "TL"
    aciklama:     Optional[str] = ""
    saha_notu:    Optional[str] = ""
    gps:          Optional[str] = ""
    durum:        Optional[str] = "Taslak"
    alanlar:      Optional[dict] = {}
    musteri_ad:   Optional[str] = ""
    musteri_tel:  Optional[str] = ""
    musteri_mail:  Optional[str] = ""
    musteri_adres: Optional[str] = ""
    musteri_tc:    Optional[str] = ""
    musteri_not:   Optional[str] = ""
    sahip_goster:  Optional[int] = 0

class IstekGiren(BaseModel):
    ad_soyad:   str
    telefon:    Optional[str] = ""
    email:      Optional[str] = ""
    mesaj:      Optional[str] = ""
    portfoy_id: Optional[int] = None

class KullaniciGiren(BaseModel):
    ad_soyad: str
    email:    str
    sifre:    str
    rol:      Optional[str] = "kullanici"

class AyarGiren(BaseModel):
    ayarlar: dict

class BlogGiren(BaseModel):
    baslik:      str
    icerik:      Optional[str] = ""
    ozet:        Optional[str] = ""
    etiketler:   Optional[list] = []
    kapak_resim: Optional[str] = ""
    durum:       Optional[str] = "Taslak"

def slug_olustur(baslik: str) -> str:
    tr_map = str.maketrans("gusiocGUSIOC", "gusiocGUSIOC")
    harfler = {"ğ":"g","ü":"u","ş":"s","ı":"i","ö":"o","ç":"c",
               "Ğ":"G","Ü":"U","Ş":"S","İ":"I","Ö":"O","Ç":"C"}
    s = baslik.lower()
    for k, v in harfler.items():
        s = s.replace(k, v)
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return s[:80]

# ─── FastAPI App ───────────────────────────────────────────────────────────────
app = FastAPI(title="Portföy Gayrimenkul API", version="1.0")
# CORS — production'da sadece kendi domain'e izin ver
IZIN_VERILEN_ORIGINLER = [
    "http://localhost:8000",
    "http://127.0.0.1:8000",
    "https://portfoygayrimenkul.com.tr",
    "https://www.portfoygayrimenkul.com.tr",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=IZIN_VERILEN_ORIGINLER,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE"],
    allow_headers=["Authorization", "Content-Type"],
)

# Güvenlik başlıkları middleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request as StarletteRequest
from starlette.responses import Response

class GuvenlikBasliklari(BaseHTTPMiddleware):
    async def dispatch(self, request: StarletteRequest, call_next):
        response: Response = await call_next(request)
        # Sadece API ve HTML yanıtlarına ekle, statik dosyalara ekleme
        if not request.url.path.startswith("/static/"):
            response.headers["X-Content-Type-Options"] = "nosniff"
            response.headers["X-Frame-Options"] = "SAMEORIGIN"
            response.headers["X-XSS-Protection"] = "1; mode=block"
            response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
            if request.url.path.startswith("/api/"):
                response.headers["Cache-Control"] = "no-store"
        return response

app.add_middleware(GuvenlikBasliklari)

# ─── Auth Endpoint'leri ────────────────────────────────────────────────────────
@app.post("/api/auth/giris")
def giris(form: OAuth2PasswordRequestForm = Depends(), request: Request = None):
    # IP al (proxy arkasında X-Forwarded-For, yoksa direkt)
    if request:
        ip = request.headers.get("X-Forwarded-For", "").split(",")[0].strip() or              getattr(request.client, "host", "bilinmiyor")
    else:
        ip = "bilinmiyor"

    # Rate limit kontrol
    rate_limit_kontrol(ip)

    # Email uzunluk kontrolü
    if not form.username or len(form.username) > 120:
        raise HTTPException(status_code=400, detail="Geçersiz istek")

    conn = get_db()
    kullanici = conn.execute(
        "SELECT * FROM kullanicilar WHERE email=? AND aktif=1",
        (form.username.lower().strip(),)
    ).fetchone()
    conn.close()

    if not kullanici or not sifre_dogrula(form.password, kullanici["sifre"]):
        rate_limit_basarisiz(ip)
        kalan = MAX_DENEME - len(_giris_denemeleri.get(ip, []))
        raise HTTPException(
            status_code=400,
            detail=f"Email veya şifre hatalı. ({max(0,kalan)} deneme hakkı kaldı)"
        )

    # Onay bekleyen kullanıcı kontrolü
    # Admin panelinden eklenen kullanıcılar (onay=1) direkt girebilir
    # Kendisi kayıt olanlar (onay=0) admin onayı bekler
    if not (kullanici["onay"] if "onay" in kullanici.keys() else 1):
        raise HTTPException(
            status_code=403,
            detail="Hesabınız henüz admin onayı bekliyor. Onaylandıktan sonra giriş yapabilirsiniz."
        )

    rate_limit_basarili(ip)
    token = token_olustur({"sub": kullanici["email"],
                            "rol": kullanici["rol"],
                            "ad":  kullanici["ad_soyad"]})
    return {"access_token": token, "token_type": "bearer",
            "rol": kullanici["rol"], "ad": kullanici["ad_soyad"]}

@app.get("/api/auth/ben")
def ben(payload=Depends(token_coz)):
    if not payload:
        return {"giris": False}
    conn = get_db()
    k = conn.execute(
        "SELECT id,ad_soyad,email,rol,COALESCE(onay,1) as onay,profil_resmi FROM kullanicilar WHERE email=?",
        (payload["sub"],)
    ).fetchone()
    conn.close()
    if not k:
        return {"giris": False}
    return {"giris": True, **dict(k)}

# ─── Portföy Endpoint'leri ─────────────────────────────────────────────────────
@app.get("/api/portfoyler")
def portfoy_listele(
    kategori: str = "",
    alt_kat:  str = "",
    durum:    str = "Aktif",
    arama:    str = "",
    payload   = Depends(token_coz)
):
    conn = get_db()
    q    = "SELECT * FROM portfoyler WHERE 1=1"
    args = []

    # Admin değilse sadece Aktif ilanlar
    is_admin = payload and payload.get("rol") == "admin"
    if not is_admin:
        q += " AND durum='Aktif'"
    elif durum:
        q += " AND durum=?"
        args.append(durum)

    if kategori:
        q += " AND ana_kategori=?"; args.append(kategori)
    if alt_kat:
        q += " AND alt_kategori=?"; args.append(alt_kat)
    if arama:
        q += " AND (baslik LIKE ? OR mahalle LIKE ? OR ilce LIKE ?)"
        args += [f"%{arama}%"] * 3

    q += " ORDER BY guncelleme DESC"
    rows = conn.execute(q, args).fetchall()
    conn.close()

    result = []
    for r in rows:
        d = dict(r)
        d["alanlar"]  = json.loads(d.get("alanlar") or "{}")
        d["resimler"] = json.loads(d.get("resimler") or "[]")
        result.append(d)
    return result

@app.get("/api/portfoyler/{pid}")
def portfoy_detay(pid: int, payload=Depends(token_coz)):
    conn = get_db()
    row = conn.execute("SELECT * FROM portfoyler WHERE id=?", (pid,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "Portföy bulunamadı")
    is_admin  = payload and payload.get("rol") == "admin"
    # Onaylı kullanıcı kontrolü
    is_onaylanmis = False
    if payload and payload.get("rol") != "admin":
        conn2 = get_db()
        k = conn2.execute("SELECT onay FROM kullanicilar WHERE email=?", (payload["sub"],)).fetchone()
        conn2.close()
        is_onaylanmis = bool(k and k["onay"])
    if not is_admin and row["durum"] != "Aktif":
        raise HTTPException(403, "Bu portföy henüz yayında değil")
    d = dict(row)
    d["alanlar"]  = json.loads(d.get("alanlar") or "{}")
    d["resimler"] = json.loads(d.get("resimler") or "[]")
    # Mal sahibi bilgileri: admin veya onaylı kullanıcı ve sahip_goster=1 ise
    yetkili = is_admin or is_onaylanmis
    if not yetkili or (not is_admin and not d.get("sahip_goster")):
        # Mal sahibi bilgilerini maskele
        d["musteri_ad"]   = "" if not yetkili else d.get("musteri_ad","")
        d["musteri_tel"]  = "" if not yetkili else d.get("musteri_tel","")
        d["musteri_mail"] = "" if not yetkili else d.get("musteri_mail","")
        d["saha_notu"]    = ""  # Saha notu her zaman sadece admin
    if not is_admin:
        d["saha_notu"] = ""
    return d

@app.post("/api/portfoyler")
def portfoy_ekle(p: PortfoyGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    c = conn.cursor()
    c.execute("""
        INSERT INTO portfoyler
        (baslik,ana_kategori,alt_kategori,ilan_tipi,il,ilce,mahalle,
         fiyat,para_birimi,aciklama,saha_notu,gps,durum,alanlar,
         musteri_ad,musteri_tel,musteri_mail,musteri_not,sahip_goster,kaynak)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,'web')
    """, (p.baslik, p.ana_kategori, p.alt_kategori, p.ilan_tipi,
          p.il, p.ilce, p.mahalle, p.fiyat, p.para_birimi,
          p.aciklama, p.saha_notu, p.gps, p.durum,
          json.dumps(p.alanlar, ensure_ascii=False),
          p.musteri_ad, p.musteri_tel, p.musteri_mail,
          p.musteri_not, p.sahip_goster))
    pid = c.lastrowid
    conn.commit(); conn.close()
    return {"id": pid, "mesaj": "Portföy oluşturuldu"}

@app.put("/api/portfoyler/{pid}")
def portfoy_guncelle(pid: int, p: PortfoyGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("""
        UPDATE portfoyler SET
        baslik=?,ana_kategori=?,alt_kategori=?,ilan_tipi=?,
        il=?,ilce=?,mahalle=?,fiyat=?,para_birimi=?,
        aciklama=?,saha_notu=?,gps=?,durum=?,alanlar=?,
        musteri_ad=?,musteri_tel=?,musteri_mail=?,musteri_not=?,sahip_goster=?,
        guncelleme=datetime('now')
        WHERE id=?
    """, (p.baslik, p.ana_kategori, p.alt_kategori, p.ilan_tipi,
          p.il, p.ilce, p.mahalle, p.fiyat, p.para_birimi,
          p.aciklama, p.saha_notu, p.gps, p.durum,
          json.dumps(p.alanlar, ensure_ascii=False),
          p.musteri_ad, p.musteri_tel, p.musteri_mail,
          p.musteri_not, p.sahip_goster, pid))
    conn.commit(); conn.close()
    return {"mesaj": "Portföy güncellendi"}

@app.patch("/api/portfoyler/{pid}/durum")
def portfoy_durum(pid: int, durum: str, payload=Depends(admin_gerek)):
    if durum not in ("Aktif", "Taslak", "Pasif", "Satıldı", "Kiralandı"):
        raise HTTPException(400, "Geçersiz durum")
    conn = get_db()
    conn.execute("UPDATE portfoyler SET durum=?,guncelleme=datetime('now') WHERE id=?",
                 (durum, pid))
    conn.commit(); conn.close()
    return {"mesaj": f"Durum → {durum}"}

@app.delete("/api/portfoyler/{pid}")
def portfoy_sil(pid: int, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("DELETE FROM portfoyler WHERE id=?", (pid,))
    conn.commit(); conn.close()
    return {"mesaj": "Portföy silindi"}

# ─── Resim Yükleme ─────────────────────────────────────────────────────────────
@app.post("/api/portfoyler/{pid}/resim")
async def resim_yukle(pid: int, dosya: UploadFile = File(...),
                      payload=Depends(admin_gerek)):
    uzanti = Path(dosya.filename or "").suffix.lower()
    if uzanti not in (".jpg", ".jpeg", ".png", ".webp"):
        raise HTTPException(400, "Sadece jpg/png/webp kabul edilir")

    # Dosyayı oku ve boyut + tip kontrol et
    icerik = await dosya.read()
    if len(icerik) > 10 * 1024 * 1024:
        raise HTTPException(400, "Dosya 10MB'dan küçük olmalı")
    # Magic byte kontrolu (JPEG, PNG, WEBP)
    jpeg_ok = len(icerik) >= 3 and icerik[0] == 0xFF and icerik[1] == 0xD8 and icerik[2] == 0xFF
    png_ok  = len(icerik) >= 4 and icerik[:4] == bytes([0x89, 0x50, 0x4E, 0x47])
    webp_ok = len(icerik) >= 12 and icerik[:4] == b'RIFF' and icerik[8:12] == b'WEBP'
    if not (jpeg_ok or png_ok or webp_ok):
        raise HTTPException(400, 'Gecersiz resim dosyasi')

    ad = f"{pid}_{uuid.uuid4().hex[:8]}{uzanti}"
    hedef = UPLOAD_DIR / ad
    with open(hedef, "wb") as f:
        f.write(icerik)

    # Portföy resimler listesine ekle
    conn = get_db()
    row = conn.execute("SELECT resimler FROM portfoyler WHERE id=?", (pid,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "Portföy bulunamadı")
    resimler = json.loads(row["resimler"] or "[]")
    url = f"/static/uploads/{ad}"
    resimler.append(url)
    conn.execute("UPDATE portfoyler SET resimler=? WHERE id=?",
                 (json.dumps(resimler), pid))
    conn.commit(); conn.close()
    return {"url": url, "resimler": resimler}

@app.delete("/api/portfoyler/{pid}/resim")
def resim_sil(pid: int, url: str, payload=Depends(admin_gerek)):
    conn = get_db()
    row = conn.execute("SELECT resimler FROM portfoyler WHERE id=?", (pid,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404)
    resimler = json.loads(row["resimler"] or "[]")
    resimler = [r for r in resimler if r != url]
    conn.execute("UPDATE portfoyler SET resimler=? WHERE id=?",
                 (json.dumps(resimler), pid))
    conn.commit(); conn.close()
    # Dosyayı da sil
    dosya = BASE_DIR / url.lstrip("/")
    if dosya.exists():
        dosya.unlink()
    return {"resimler": resimler}

# ─── Resim Sıralama ────────────────────────────────────────────────────────────
class ResimSirala(BaseModel):
    resimler: list

@app.put("/api/portfoyler/{pid}/resim/sirala")
def resim_sirala(pid: int, data: ResimSirala, payload=Depends(admin_gerek)):
    conn = get_db()
    row = conn.execute("SELECT resimler FROM portfoyler WHERE id=?", (pid,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404, "Portföy bulunamadı")
    mevcut = set(json.loads(row["resimler"] or "[]"))
    yeni   = [r for r in data.resimler if r in mevcut]
    conn.execute("UPDATE portfoyler SET resimler=?, guncelleme=datetime('now') WHERE id=?",
                 (json.dumps(yeni), pid))
    conn.commit(); conn.close()
    return {"resimler": yeni}

@app.patch("/api/portfoyler/{pid}/resim/kapak")
def kapak_sec(pid: int, url: str, payload=Depends(admin_gerek)):
    conn = get_db()
    row = conn.execute("SELECT resimler FROM portfoyler WHERE id=?", (pid,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(404)
    resimler = json.loads(row["resimler"] or "[]")
    if url not in resimler:
        conn.close()
        raise HTTPException(400, "Resim bulunamadı")
    resimler = [url] + [r for r in resimler if r != url]
    conn.execute("UPDATE portfoyler SET resimler=? WHERE id=?", (json.dumps(resimler), pid))
    conn.commit(); conn.close()
    return {"resimler": resimler}

# ─── Şifre & Profil ────────────────────────────────────────────────────────────
class SifreDegistir(BaseModel):
    mevcut_sifre: str
    yeni_sifre:   str

# ── Şifre Sıfırlama (Admin için terminal token) ──────────────────────────────
_sifre_sifirlama_tokenlar: dict = {}  # token → {email, exp}

@app.post("/api/auth/sifre-sifirlama-baslat")
def sifre_sifirlama_baslat(email_data: dict):
    """Admin şifresini sıfırlamak için tek kullanımlık token üretir.
    Token terminale yazdırılır — dışarıya e-posta gönderilmez (güvenli)."""
    email = email_data.get("email", "").lower().strip()
    conn = get_db()
    k = conn.execute(
        "SELECT id,email,rol FROM kullanicilar WHERE email=? AND aktif=1", (email,)
    ).fetchone()
    conn.close()
    # Kullanıcı var/yok bilgisi verme (güvenlik)
    token = secrets.token_urlsafe(32)
    exp = time.time() + 900  # 15 dakika
    _sifre_sifirlama_tokenlar[token] = {"email": email, "exp": exp}
    # Terminale yaz — sadece sunucu sahibi görebilir
    print(f"\n{'='*60}")
    print(f"ŞİFRE SIFIRLAMA TOKENI")
    print(f"E-posta : {email}")
    print(f"Token   : {token}")
    print(f"Geçerlilik: 15 dakika")
    print(f"{'='*60}\n")
    return {"mesaj": "Sıfırlama talebi alındı. Sunucu terminalini kontrol edin."}

@app.post("/api/auth/sifre-sifirlama-tamamla")
def sifre_sifirlama_tamamla(data: dict):
    """Token + yeni şifre ile şifreyi sıfırla."""
    token = data.get("token", "").strip()
    yeni = data.get("yeni_sifre", "")
    if not token or not yeni:
        raise HTTPException(400, "Token ve yeni şifre gerekli")
    if len(yeni) < 8:
        raise HTTPException(400, "Şifre en az 8 karakter olmalı")
    info = _sifre_sifirlama_tokenlar.get(token)
    if not info:
        raise HTTPException(400, "Geçersiz veya süresi dolmuş token")
    if time.time() > info["exp"]:
        del _sifre_sifirlama_tokenlar[token]
        raise HTTPException(400, "Token süresi dolmuş (15 dakika)")
    conn = get_db()
    conn.execute(
        "UPDATE kullanicilar SET sifre=? WHERE email=?",
        (hash_sifre(yeni), info["email"])
    )
    conn.commit(); conn.close()
    del _sifre_sifirlama_tokenlar[token]
    return {"mesaj": "Şifre başarıyla güncellendi"}

@app.put("/api/kullanicilar/sifre")
def sifre_degistir(data: SifreDegistir, payload=Depends(kullanici_gerek)):
    if len(data.yeni_sifre) < 8:
        raise HTTPException(400, "Yeni şifre en az 8 karakter olmalı")
    conn = get_db()
    k = conn.execute("SELECT * FROM kullanicilar WHERE email=?", (payload["sub"],)).fetchone()
    if not k or not sifre_dogrula(data.mevcut_sifre, k["sifre"]):
        conn.close()
        raise HTTPException(400, "Mevcut şifre hatalı")
    conn.execute("UPDATE kullanicilar SET sifre=? WHERE email=?",
                 (hash_sifre(data.yeni_sifre), payload["sub"]))
    conn.commit(); conn.close()
    return {"mesaj": "Şifre başarıyla güncellendi"}

class ProfilGuncelle(BaseModel):
    ad_soyad: str
    email:    str

@app.put("/api/kullanicilar/profil")
def profil_guncelle(data: ProfilGuncelle, payload=Depends(kullanici_gerek)):
    if not data.ad_soyad or not data.email:
        raise HTTPException(400, "Ad soyad ve e-posta zorunlu")
    conn = get_db()
    mevcut = conn.execute("SELECT id FROM kullanicilar WHERE email=? AND email!=?",
                          (data.email, payload["sub"])).fetchone()
    if mevcut:
        conn.close()
        raise HTTPException(400, "Bu e-posta başka bir hesapta kayıtlı")
    conn.execute("UPDATE kullanicilar SET ad_soyad=?, email=? WHERE email=?",
                 (data.ad_soyad, data.email, payload["sub"]))
    conn.commit(); conn.close()
    return {"mesaj": "Profil güncellendi"}

# ─── Belge Parser Endpoint'i ────────────────────────────────────────────────────
@app.post("/api/belge/parse")
async def belge_parse(dosya: UploadFile = File(...), payload=Depends(admin_gerek)):
    gecici = BASE_DIR / "static" / "uploads" / f"tmp_{uuid.uuid4().hex}"
    uzanti = Path(dosya.filename).suffix.lower()

    with open(str(gecici) + uzanti, "wb") as f:
        shutil.copyfileobj(dosya.file, f)

    try:
        if uzanti == ".docx":
            sonuc = docx_parse(str(gecici) + uzanti)
        elif uzanti in (".html", ".htm"):
            icerik = open(str(gecici) + uzanti, encoding="utf-8", errors="ignore").read()
            sonuc = html_parse(icerik)
        else:
            raise HTTPException(400, "Sadece .docx veya .html destekleniyor")
    finally:
        try:
            Path(str(gecici) + uzanti).unlink()
        except:
            pass

    # İlgili alan şablonunu da döndür
    alan_list = alan_sablonu_sec(
        sonuc.get("ana_kategori", "Konut"),
        sonuc.get("alt_kategori", "Satılık"),
        sonuc.get("ilan_tipi", "")
    )
    return {"portfoy": sonuc, "alan_sablonu": alan_list}

# ─── Kategori Endpoint'i ───────────────────────────────────────────────────────
@app.get("/api/kategoriler")
def kategoriler():
    return {
        "kategoriler": KATEGORILER,
        "ilan_tipleri": ILAN_TIPLERI
    }

@app.get("/api/alanlar")
def alan_sec(ana_kat: str, alt_kat: str, ilan_tipi: str = ""):
    return alan_sablonu_sec(ana_kat, alt_kat, ilan_tipi)

# ─── Ziyaretçi İstek Formu ────────────────────────────────────────────────────
@app.post("/api/istekler")
def istek_gonder(istek: IstekGiren):
    conn = get_db()
    conn.execute("""INSERT INTO kullanici_istekleri
                    (ad_soyad,telefon,email,mesaj,portfoy_id)
                    VALUES (?,?,?,?,?)""",
                 (istek.ad_soyad, istek.telefon, istek.email,
                  istek.mesaj, istek.portfoy_id))
    conn.commit(); conn.close()
    return {"mesaj": "İsteğiniz alındı, en kısa sürede dönüş yapılacak."}

@app.get("/api/istekler")
def istek_listele(payload=Depends(admin_gerek)):
    conn = get_db()
    rows = conn.execute("""
        SELECT i.*, p.baslik as portfoy_baslik
        FROM kullanici_istekleri i
        LEFT JOIN portfoyler p ON i.portfoy_id=p.id
        ORDER BY i.olusturma DESC
    """).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.patch("/api/istekler/{iid}/durum")
def istek_durum(iid: int, durum: str, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("UPDATE kullanici_istekleri SET durum=? WHERE id=?", (durum, iid))
    conn.commit(); conn.close()
    return {"mesaj": "Güncellendi"}

# ─── Kullanıcı Yönetimi ────────────────────────────────────────────────────────
@app.patch("/api/kullanicilar/{kid}/onay-kaldir")
def kullanici_onay_kaldir(kid: int, payload=Depends(admin_gerek)):
    if kid == 1:
        raise HTTPException(400, "Varsayılan admin değiştirilemez")
    conn = get_db()
    conn.execute("UPDATE kullanicilar SET onay=0 WHERE id=?", (kid,))
    conn.commit(); conn.close()
    return {"mesaj": "Kullanıcı onayı kaldırıldı"}

@app.get("/api/kullanicilar")
def kullanici_listele(payload=Depends(admin_gerek)):
    conn = get_db()
    rows = conn.execute(
        "SELECT id,ad_soyad,email,rol,aktif,COALESCE(onay,1) as onay,COALESCE(onay,1) as onayli,profil_resmi,olusturma FROM kullanicilar ORDER BY id"
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/api/kullanicilar/kayit")
def kullanici_kayit(k: KullaniciGiren):
    """Ziyaretçi kayıt — admin onayı gerekir, onay=0 ile oluşturulur."""
    if not k.email or not k.sifre or len(k.sifre) < 6:
        raise HTTPException(400, "Geçerli email ve en az 6 karakterli şifre gerekli")
    conn = get_db()
    try:
        conn.execute(
            "INSERT INTO kullanicilar (ad_soyad,email,sifre,rol,onay) VALUES (?,?,?,?,?)",
            (k.ad_soyad, k.email.lower().strip(), hash_sifre(k.sifre), "kullanici", 0)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        raise HTTPException(400, "Bu e-posta adresi zaten kayıtlı")
    finally:
        conn.close()
    return {"mesaj": "Kayıt alındı. Admin onayından sonra giriş yapabilirsiniz."}

@app.post("/api/kullanicilar")
def kullanici_ekle(k: KullaniciGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    try:
        # Admin tarafından eklenen kullanıcılar direkt onaylı, kayıt olursa 0
        onay_durum = 1  # Admin panelinden eklenince otomatik onaylı
        conn.execute(
            "INSERT INTO kullanicilar (ad_soyad,email,sifre,rol,onay) VALUES (?,?,?,?,?)",
            (k.ad_soyad, k.email, hash_sifre(k.sifre), k.rol, onay_durum)
        )
        conn.commit()
    except sqlite3.IntegrityError:
        raise HTTPException(400, "Bu email zaten kayıtlı")
    finally:
        conn.close()
    return {"mesaj": "Kullanıcı oluşturuldu"}

@app.patch("/api/kullanicilar/{kid}/onayla")
def kullanici_onayla(kid: int, payload=Depends(admin_gerek)):
    conn = get_db()
    k = conn.execute("SELECT onay FROM kullanicilar WHERE id=?", (kid,)).fetchone()
    if not k:
        conn.close()
        raise HTTPException(404, "Kullanıcı bulunamadı")
    yeni_onay = 1  # Onayla endpoint'i her zaman onaylar
    conn.execute("UPDATE kullanicilar SET onay=? WHERE id=?", (yeni_onay, kid))
    conn.commit(); conn.close()
    return {"mesaj": "Onaylı" if yeni_onay else "Onay kaldırıldı", "onay": yeni_onay}

@app.delete("/api/kullanicilar/{kid}")
def kullanici_sil(kid: int, payload=Depends(admin_gerek)):
    if kid == 1:
        raise HTTPException(400, "Varsayılan admin silinemez")
    conn = get_db()
    conn.execute("DELETE FROM kullanicilar WHERE id=?", (kid,))
    conn.commit(); conn.close()
    return {"mesaj": "Kullanıcı silindi"}


# ─── Logo Yükleme ──────────────────────────────────────────────────────────────
@app.post("/api/kullanicilar/profil-resmi")
async def profil_resmi_yukle(dosya: UploadFile = File(...), payload=Depends(kullanici_gerek)):
    uzanti = Path(dosya.filename or "").suffix.lower()
    if uzanti not in (".jpg", ".jpeg", ".png", ".webp"):
        raise HTTPException(400, "Sadece jpg/png/webp kabul edilir")

    icerik = await dosya.read()
    if len(icerik) > 5 * 1024 * 1024:
        raise HTTPException(400, "Profil resmi 5MB'dan küçük olmalı")

    # Otomatik boyutlandırma: kare kırp + 400x400 WebP
    ad = f"profil_{uuid.uuid4().hex[:10]}.webp"
    hedef = UPLOAD_DIR / ad
    try:
        import io as _io
        from PIL import Image, ImageOps
        img = Image.open(_io.BytesIO(icerik))
        img = ImageOps.exif_transpose(img)  # EXIF rotasyon düzelt
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        # Kare kırp — merkez odaklı
        min_kenar = min(img.width, img.height)
        sol = (img.width  - min_kenar) // 2
        ust = (img.height - min_kenar) // 2
        img = img.crop((sol, ust, sol + min_kenar, ust + min_kenar))
        img = img.resize((400, 400), Image.LANCZOS)
        img.save(hedef, "WEBP", quality=88, method=4)
    except Exception:
        # Pillow başarısız: orijinali kaydet
        with open(hedef, "wb") as fw:
            fw.write(icerik)

    url = f"/static/uploads/{ad}"
    conn = get_db()
    # Eski resmi sil
    eski = conn.execute("SELECT profil_resmi FROM kullanicilar WHERE email=?", (payload["sub"],)).fetchone()
    if eski and eski["profil_resmi"]:
        eski_dosya = BASE_DIR / eski["profil_resmi"].lstrip("/")
        if eski_dosya.exists():
            try: eski_dosya.unlink()
            except: pass
    conn.execute("UPDATE kullanicilar SET profil_resmi=? WHERE email=?", (url, payload["sub"]))
    conn.commit(); conn.close()
    return {"url": url}

@app.get("/api/kullanicilar/ben")
def kullanici_ben(payload=Depends(kullanici_gerek)):
    conn = get_db()
    k = conn.execute("SELECT id,ad_soyad,email,rol,aktif,onay,profil_resmi FROM kullanicilar WHERE email=?",
                     (payload["sub"],)).fetchone()
    conn.close()
    if not k: raise HTTPException(404, "Kullanıcı bulunamadı")
    return dict(k)

@app.get("/api/portfoyler/{pid}/danismanlar")
def portfoy_danismanlar(pid: int):
    """Portföy sahibinin danışman bilgileri (isim + profil resmi)"""
    conn = get_db()
    p = conn.execute("SELECT musteri_ad, musteri_tel FROM portfoyler WHERE id=?", (pid,)).fetchone()
    conn.close()
    if not p: raise HTTPException(404)
    return {"musteri_ad": p["musteri_ad"], "musteri_tel": p["musteri_tel"]}

@app.post("/api/logo/yukle")
async def logo_yukle(dosya: UploadFile = File(...), payload=Depends(admin_gerek)):
    uzanti = Path(dosya.filename).suffix.lower()
    if uzanti not in (".jpg", ".jpeg", ".png", ".webp", ".svg"):
        raise HTTPException(400, "Desteklenen formatlar: jpg, png, webp, svg")
    # Eski logoyu sil
    conn = get_db()
    eski = conn.execute("SELECT deger FROM site_ayarlari WHERE anahtar='logo_url'").fetchone()
    if eski and eski["deger"]:
        eski_dosya = BASE_DIR / eski["deger"].lstrip("/")
        if eski_dosya.exists():
            eski_dosya.unlink()
    conn.close()
    # Yeni logo kaydet
    ad = f"logo{uzanti}"
    hedef = BASE_DIR / "static" / "img" / ad
    hedef.parent.mkdir(parents=True, exist_ok=True)
    with open(hedef, "wb") as f:
        shutil.copyfileobj(dosya.file, f)
    url = f"/static/img/{ad}"
    conn = get_db()
    conn.execute("INSERT OR REPLACE INTO site_ayarlari VALUES ('logo_url',?)", (url,))
    conn.commit(); conn.close()
    return {"url": url, "mesaj": "Logo yüklendi"}

@app.delete("/api/logo")
def logo_sil(payload=Depends(admin_gerek)):
    conn = get_db()
    row = conn.execute("SELECT deger FROM site_ayarlari WHERE anahtar='logo_url'").fetchone()
    if row and row["deger"]:
        dosya = BASE_DIR / row["deger"].lstrip("/")
        if dosya.exists():
            dosya.unlink()
    conn.execute("INSERT OR REPLACE INTO site_ayarlari VALUES ('logo_url','')")
    conn.commit(); conn.close()
    return {"mesaj": "Logo silindi"}

# ─── Site Ayarları ─────────────────────────────────────────────────────────────
@app.get("/api/ayarlar")
def ayar_getir():
    conn = get_db()
    rows = conn.execute("SELECT * FROM site_ayarlari").fetchall()
    conn.close()
    return {r["anahtar"]: r["deger"] for r in rows}

@app.put("/api/ayarlar")
def ayar_kaydet(data: AyarGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    for k, v in data.ayarlar.items():
        conn.execute("INSERT OR REPLACE INTO site_ayarlari VALUES (?,?)", (k, str(v)))
    conn.commit(); conn.close()
    return {"mesaj": "Ayarlar kaydedildi"}

# ─── İstatistik ────────────────────────────────────────────────────────────────
@app.get("/api/istatistik")
def istatistik(payload=Depends(token_coz)):
    conn = get_db()
    aktif = conn.execute("SELECT COUNT(*) FROM portfoyler WHERE durum='Aktif'").fetchone()[0]
    kat_dag = conn.execute("""
        SELECT ana_kategori, COUNT(*) as sayi
        FROM portfoyler WHERE durum='Aktif'
        GROUP BY ana_kategori
    """).fetchall()
    sonuc = {
        "toplam": aktif,
        "aktif":  aktif,
        "taslak": 0,
        "yeni_istekler": 0,
        "kategori_dagilimi": [dict(r) for r in kat_dag],
    }
    # Admin ise ek bilgiler
    if payload and payload.get("rol") == "admin":
        toplam   = conn.execute("SELECT COUNT(*) FROM portfoyler").fetchone()[0]
        taslak   = conn.execute("SELECT COUNT(*) FROM portfoyler WHERE durum='Taslak'").fetchone()[0]
        istekler = conn.execute("SELECT COUNT(*) FROM kullanici_istekleri WHERE durum='Yeni'").fetchone()[0]
        sonuc.update({"toplam": toplam, "taslak": taslak, "yeni_istekler": istekler})
    conn.close()
    return sonuc

# ─── AI Fiyat Analizi ──────────────────────────────────────────────────────────
class AIAyarGiren(BaseModel):
    ai_api_key: str = ""
    ai_saglayici: str = "deepseek"

AI_SAGLAYICI_URLS = {
    "deepseek": "https://api.deepseek.com/v1/chat/completions",
    "groq":     "https://api.groq.com/openai/v1/chat/completions",
    "openai":   "https://api.openai.com/v1/chat/completions",
}
AI_SAGLAYICI_MODEL = {
    "deepseek": "deepseek-chat",
    "groq":     "llama-3.3-70b-versatile",
    "openai":   "gpt-4o-mini",
}

@app.put("/api/ayarlar/ai")
def ai_ayar_kaydet(data: AIAyarGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("INSERT OR REPLACE INTO site_ayarlari VALUES ('ai_api_key',?)", (data.ai_api_key,))
    conn.execute("INSERT OR REPLACE INTO site_ayarlari VALUES ('ai_saglayici',?)", (data.ai_saglayici,))
    conn.commit(); conn.close()
    return {"mesaj": "AI ayarları kaydedildi"}

def _m2_cikar(alanlar: dict) -> float:
    try:
        return float(alanlar.get("net_m2") or alanlar.get("alan_m2") or 0)
    except (ValueError, TypeError):
        return 0.0

def _fiyat_sayi(fiyat_str: str) -> float:
    try:
        temiz = re.sub(r"[^\d.,]", "", fiyat_str or "")
        temiz = temiz.replace(".", "").replace(",", ".")
        return float(temiz) if temiz else 0.0
    except (ValueError, TypeError):
        return 0.0

@app.get("/api/portfoyler/{pid}/fiyat-analizi")
async def fiyat_analizi(pid: int, payload=Depends(admin_gerek)):
    conn = get_db()
    hedef = conn.execute("SELECT * FROM portfoyler WHERE id=?", (pid,)).fetchone()
    if not hedef:
        conn.close()
        raise HTTPException(404, "Portföy bulunamadı")

    hedef_alanlar = json.loads(hedef["alanlar"] or "{}")
    hedef_m2 = _m2_cikar(hedef_alanlar)
    hedef_fiyat = _fiyat_sayi(hedef["fiyat"])

    if hedef_m2 <= 0 or hedef_fiyat <= 0:
        conn.close()
        return {
            "yeterli_veri": False,
            "mesaj": "Bu ilanda m² veya fiyat bilgisi eksik, analiz yapılamıyor."
        }

    benzerler = conn.execute("""
        SELECT * FROM portfoyler
        WHERE id != ? AND durum='Aktif'
        AND ana_kategori=? AND alt_kategori=?
    """, (pid, hedef["ana_kategori"], hedef["alt_kategori"])).fetchall()
    conn.close()

    karsilastirma = []
    for b in benzerler:
        b_alanlar = json.loads(b["alanlar"] or "{}")
        b_m2 = _m2_cikar(b_alanlar)
        b_fiyat = _fiyat_sayi(b["fiyat"])
        if b_m2 > 0 and b_fiyat > 0:
            karsilastirma.append({
                "id": b["id"], "baslik": b["baslik"], "mahalle": b["mahalle"],
                "m2": b_m2, "fiyat": b_fiyat, "m2_fiyat": round(b_fiyat / b_m2, 2)
            })

    if len(karsilastirma) < 3:
        return {
            "yeterli_veri": False,
            "mesaj": f"Karşılaştırma için yeterli veri yok ({len(karsilastirma)} benzer ilan bulundu, minimum 3 gerekli).",
            "benzer_sayisi": len(karsilastirma)
        }

    m2_fiyatlari = sorted([k["m2_fiyat"] for k in karsilastirma])
    ortalama_m2_fiyat = sum(m2_fiyatlari) / len(m2_fiyatlari)
    medyan_m2_fiyat = m2_fiyatlari[len(m2_fiyatlari)//2]
    hedef_m2_fiyat = round(hedef_fiyat / hedef_m2, 2)
    fark_yuzde = round(((hedef_m2_fiyat - ortalama_m2_fiyat) / ortalama_m2_fiyat) * 100, 1)

    if fark_yuzde > 15:
        durum = "yuksek"; durum_metin = "Piyasa Ortalamasının Üzerinde"
    elif fark_yuzde < -15:
        durum = "dusuk"; durum_metin = "Piyasa Ortalamasının Altında"
    else:
        durum = "uygun"; durum_metin = "Piyasa Ortalamasına Yakın"

    sonuc = {
        "yeterli_veri": True,
        "hedef_m2_fiyat": hedef_m2_fiyat,
        "ortalama_m2_fiyat": round(ortalama_m2_fiyat, 2),
        "medyan_m2_fiyat": round(medyan_m2_fiyat, 2),
        "fark_yuzde": fark_yuzde,
        "durum": durum,
        "durum_metin": durum_metin,
        "benzer_sayisi": len(karsilastirma),
        "en_yakin_3": sorted(karsilastirma, key=lambda x: abs(x["m2_fiyat"] - hedef_m2_fiyat))[:3],
        "ai_yorum": None,
    }

    conn2 = get_db()
    ai_key_row = conn2.execute("SELECT deger FROM site_ayarlari WHERE anahtar='ai_api_key'").fetchone()
    ai_saglayici_row = conn2.execute("SELECT deger FROM site_ayarlari WHERE anahtar='ai_saglayici'").fetchone()
    conn2.close()
    ai_key = ai_key_row["deger"] if ai_key_row else ""
    ai_saglayici = (ai_saglayici_row["deger"] if ai_saglayici_row else "") or "deepseek"

    if ai_key:
        try:
            import httpx
            prompt = f"""Sen bir Fethiye/Muğla gayrimenkul danışmanısın. Aşağıdaki ilan için kısa, profesyonel bir piyasa değerlendirmesi yaz (maksimum 3-4 cümle, Türkçe):

İlan: {hedef['baslik']}
Kategori: {hedef['ana_kategori']} / {hedef['alt_kategori']}
Konum: {hedef['mahalle']}, {hedef['ilce']}
m²: {hedef_m2}
Fiyat: {hedef['fiyat']} {hedef['para_birimi']}
m² birim fiyat: {hedef_m2_fiyat} TL/m²
Bölge ortalaması: {round(ortalama_m2_fiyat,2)} TL/m²
Fark: %{fark_yuzde} ({durum_metin})
Karşılaştırılan ilan sayısı: {len(karsilastirma)}

Sadece değerlendirme metnini yaz, başlık veya markdown kullanma."""

            url = AI_SAGLAYICI_URLS.get(ai_saglayici, AI_SAGLAYICI_URLS["deepseek"])
            model = AI_SAGLAYICI_MODEL.get(ai_saglayici, "deepseek-chat")
            async with httpx.AsyncClient(timeout=20.0) as client:
                r = await client.post(url,
                    headers={"Authorization": f"Bearer {ai_key}", "Content-Type": "application/json"},
                    json={"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.4, "max_tokens": 300}
                )
                if r.status_code == 200:
                    data = r.json()
                    sonuc["ai_yorum"] = data["choices"][0]["message"]["content"].strip()
        except Exception as e:
            sonuc["ai_yorum_hata"] = str(e)[:200]

    return sonuc

@app.get("/api/fiyat-analizi/genel")
def fiyat_analizi_genel(payload=Depends(admin_gerek)):
    conn = get_db()
    rows = conn.execute("SELECT ana_kategori, alt_kategori, fiyat, alanlar FROM portfoyler WHERE durum='Aktif'").fetchall()
    conn.close()

    gruplar = {}
    for r in rows:
        alanlar = json.loads(r["alanlar"] or "{}")
        m2 = _m2_cikar(alanlar)
        fiyat = _fiyat_sayi(r["fiyat"])
        if m2 <= 0 or fiyat <= 0:
            continue
        anahtar = f"{r['ana_kategori']} / {r['alt_kategori']}"
        gruplar.setdefault(anahtar, []).append(round(fiyat / m2, 2))

    sonuc = []
    for kategori, fiyatlar in gruplar.items():
        if len(fiyatlar) < 2:
            continue
        sonuc.append({
            "kategori": kategori,
            "ilan_sayisi": len(fiyatlar),
            "min_m2_fiyat": round(min(fiyatlar), 2),
            "max_m2_fiyat": round(max(fiyatlar), 2),
            "ortalama_m2_fiyat": round(sum(fiyatlar) / len(fiyatlar), 2),
        })
    return sorted(sonuc, key=lambda x: -x["ilan_sayisi"])


# ─── PDF İlan Broşürü ───────────────────────────────────────────────────────────
@app.get("/api/portfoyler/{pid}/pdf")
def pdf_brosur(pid: int, payload=Depends(admin_gerek)):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    import io

    conn = get_db()
    p = conn.execute("SELECT * FROM portfoyler WHERE id=?", (pid,)).fetchone()
    ayarlar_rows = conn.execute("SELECT * FROM site_ayarlari").fetchall()
    conn.close()
    if not p:
        raise HTTPException(404, "Portföy bulunamadı")

    ayarlar = {r["anahtar"]: r["deger"] for r in ayarlar_rows}
    alanlar = json.loads(p["alanlar"] or "{}")
    resimler = json.loads(p["resimler"] or "[]")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    W, H = A4

    KIREMIT = colors.HexColor("#C45C35")
    TOPRAK  = colors.HexColor("#2D2016")
    GRI     = colors.HexColor("#7A6E65")

    c.setFillColor(KIREMIT)
    c.rect(0, H-25*mm, W, 25*mm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(15*mm, H-15*mm, ayarlar.get("site_adi", "Portföy Gayrimenkul"))
    c.setFont("Helvetica", 9)
    c.drawString(15*mm, H-20*mm, ayarlar.get("web_sitesi", "portfoygayrimenkul.com.tr"))
    c.setFont("Helvetica", 9)
    c.drawRightString(W-15*mm, H-15*mm, ayarlar.get("telefon", ""))
    c.drawRightString(W-15*mm, H-20*mm, ayarlar.get("email", ""))

    y = H - 35*mm

    if resimler:
        try:
            img_path = BASE_DIR / resimler[0].lstrip("/")
            if img_path.exists():
                img = ImageReader(str(img_path))
                iw, ih = img.getSize()
                foto_w = W - 30*mm
                foto_h = 75*mm
                c.saveState()
                p_clip = c.beginPath()
                p_clip.rect(15*mm, y-foto_h, foto_w, foto_h)
                c.clipPath(p_clip, stroke=0)
                oran = max(foto_w/iw, foto_h/ih)
                disp_w, disp_h = iw*oran, ih*oran
                ox = 15*mm - (disp_w-foto_w)/2
                oy = (y-foto_h) - (disp_h-foto_h)/2
                c.drawImage(img, ox, oy, width=disp_w, height=disp_h)
                c.restoreState()
                y -= foto_h + 8*mm
        except Exception:
            pass

    c.setFillColor(TOPRAK)
    c.setFont("Helvetica-Bold", 17)
    c.drawString(15*mm, y, p["baslik"][:60])
    y -= 8*mm

    c.setFillColor(KIREMIT)
    c.setFont("Helvetica-Bold", 20)
    c.drawString(15*mm, y, f"{p['fiyat'] or 'Fiyat Sorunuz'} {p['para_birimi'] or ''}")
    y -= 9*mm

    c.setFillColor(GRI)
    c.setFont("Helvetica", 10)
    kat_metin = f"{p['ana_kategori']} / {p['alt_kategori']}" + (f" / {p['ilan_tipi']}" if p["ilan_tipi"] else "")
    c.drawString(15*mm, y, kat_metin)
    y -= 5*mm
    konum = " / ".join(filter(None, [p["mahalle"], p["ilce"], p["il"]]))
    c.drawString(15*mm, y, f"Konum: {konum}")
    y -= 9*mm

    c.setStrokeColor(colors.HexColor("#E8DFD0"))
    c.line(15*mm, y, W-15*mm, y)
    y -= 7*mm

    teknik = [(k.replace("_", " ").title(), v) for k, v in alanlar.items() if v and k != "ozellikler"]
    if teknik:
        c.setFillColor(TOPRAK)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(15*mm, y, "Teknik Bilgiler")
        y -= 6*mm
        c.setFont("Helvetica", 9)
        col_w = (W - 30*mm) / 2
        for i, (k, v) in enumerate(teknik[:14]):
            col = i % 2
            row = i // 2
            xx = 15*mm + col * col_w
            yy = y - row * 5.5*mm
            c.setFillColor(GRI)
            c.drawString(xx, yy, f"{k}:")
            c.setFillColor(TOPRAK)
            c.drawString(xx + 32*mm, yy, str(v)[:30])
        satir_sayisi = (min(len(teknik), 14) + 1) // 2
        y -= satir_sayisi * 5.5*mm + 6*mm

    if p["aciklama"]:
        c.setFillColor(TOPRAK)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(15*mm, y, "Açıklama")
        y -= 6*mm
        c.setFont("Helvetica", 9)
        c.setFillColor(colors.HexColor("#3D362E"))
        from textwrap import wrap
        for line in wrap(p["aciklama"], 95)[:8]:
            c.drawString(15*mm, y, line)
            y -= 4.5*mm
        y -= 4*mm

    if alanlar.get("ozellikler"):
        c.setFillColor(TOPRAK)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(15*mm, y, "Öne Çıkan Özellikler: ")
        c.setFont("Helvetica", 9)
        c.setFillColor(GRI)
        c.drawString(50*mm, y, str(alanlar["ozellikler"])[:80])
        y -= 8*mm

    c.setFillColor(colors.HexColor("#E8DFD0"))
    c.rect(0, 0, W, 18*mm, fill=1, stroke=0)
    c.setFillColor(GRI)
    c.setFont("Helvetica", 8)
    c.drawCentredString(W/2, 10*mm,
        f"{ayarlar.get('site_adi','')} · {ayarlar.get('web_sitesi','')} · {ayarlar.get('telefon','')}")
    c.drawCentredString(W/2, 6*mm, "Bu broşür otomatik oluşturulmuştur.")

    c.showPage()
    c.save()
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    dosya_adi = re.sub(r'[^a-zA-Z0-9_-]', '_', p["baslik"])[:40]
    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{dosya_adi}.pdf"'}
    )


@app.get("/api/portfoyler/{pid}/sahip-profil")
def portfoy_sahip_profil(pid: int, payload=Depends(token_coz)):
    """İlanın mal sahibine atanan kullanıcının profil resmini döner."""
    conn = get_db()
    p = conn.execute("SELECT yazar_id, musteri_ad FROM portfoyler WHERE id=?", (pid,)).fetchone()
    if not p:
        conn.close()
        raise HTTPException(404)
    # Portföyü oluşturan adminin profil resmi
    admin = conn.execute(
        "SELECT profil_resmi FROM kullanicilar WHERE rol='admin' ORDER BY id LIMIT 1"
    ).fetchone()
    conn.close()
    return {
        "profil_resmi": admin["profil_resmi"] if admin else "",
        "musteri_ad": p["musteri_ad"] or ""
    }


# ─── Banner API ────────────────────────────────────────────────────────────────

class BannerGiren(BaseModel):
    tip:           str = "slider"
    baslik:        Optional[str] = ""
    metin:         Optional[str] = ""
    link_url:      Optional[str] = ""
    link_metin:    Optional[str] = "İncele"
    konum:         Optional[str] = "anasayfa"
    boyut:         Optional[str] = "normal"
    renk_arkaplan: Optional[str] = "#C45C35"
    renk_metin:    Optional[str] = "#ffffff"
    sira:          Optional[int] = 0
    aktif:         Optional[int] = 1

@app.get("/api/bannerlar")
def banner_listele(konum: str = "", sadece_aktif: bool = False):
    conn = get_db()
    q = "SELECT * FROM bannerlar WHERE 1=1"
    args = []
    if konum:
        q += " AND konum=?"; args.append(konum)
    if sadece_aktif:
        q += " AND aktif=1"
    q += " ORDER BY sira, id"
    rows = conn.execute(q, args).fetchall()
    conn.close()
    return [dict(r) for r in rows]

# Banner Sabitleri
BANNER_KONUMLAR = {
    "anasayfa_ust":       "Ana Sayfa — En Üst",
    "anasayfa_hero_alti": "Ana Sayfa — Hero Altı",
    "ilanlar_ust":        "İlanlar Sayfası — Üst",
    "haberler_ust":       "Haberler Sayfası — Üst",
    "tum_sayfalar_ust":   "Tüm Sayfalarda — Navbar Altı",
    "tum_sayfalar_alt":   "Tüm Sayfalarda — Footer Üstü",
}
BANNER_BOYUTLAR = {
    "tam":   {"label": "Tam Genişlik",  "yukseklik": 400},
    "genis": {"label": "Geniş",         "yukseklik": 300},
    "orta":  {"label": "Orta",          "yukseklik": 220},
    "ince":  {"label": "İnce Şerit",    "yukseklik": 120},
}

@app.get("/api/bannerlar/konumlar")
def banner_konumlar():
    return {"konumlar": BANNER_KONUMLAR, "boyutlar": BANNER_BOYUTLAR}

@app.post("/api/bannerlar")
def banner_ekle(data: dict, payload=Depends(admin_gerek)):
    conn = get_db()
    c = conn.cursor()
    c.execute("""INSERT INTO bannerlar
        (baslik, alt_metin, link_url, link_hedef, tip, konum, boyut, sira, aktif)
        VALUES (?,?,?,?,?,?,?,?,?)""",
        (data.get("baslik",""), data.get("alt_metin",""),
         data.get("link_url",""), data.get("link_hedef","_self"),
         data.get("tip","slider"), data.get("konum","anasayfa_hero_alti"),
         data.get("boyut","genis"), data.get("sira",0), data.get("aktif",1)))
    bid = c.lastrowid
    conn.commit(); conn.close()
    return {"id": bid, "mesaj": "Banner oluşturuldu"}

@app.put("/api/bannerlar/{bid}")
def banner_guncelle(bid: int, data: dict, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("""UPDATE bannerlar SET
        baslik=?, alt_metin=?, link_url=?, link_hedef=?,
        tip=?, konum=?, boyut=?, sira=?, aktif=?
        WHERE id=?""",
        (data.get("baslik",""), data.get("alt_metin",""),
         data.get("link_url",""), data.get("link_hedef","_self"),
         data.get("tip","slider"), data.get("konum","anasayfa_hero_alti"),
         data.get("boyut","genis"), data.get("sira",0), data.get("aktif",1), bid))
    conn.commit(); conn.close()
    return {"mesaj": "Güncellendi"}

@app.patch("/api/bannerlar/{bid}/aktif")
def banner_aktif_toggle(bid: int, aktif: int, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("UPDATE bannerlar SET aktif=? WHERE id=?", (aktif, bid))
    conn.commit(); conn.close()
    return {"mesaj": "Durum güncellendi"}

@app.patch("/api/bannerlar/sira")
def banner_sira(siralar: list, payload=Depends(admin_gerek)):
    conn = get_db()
    for i, bid in enumerate(siralar):
        conn.execute("UPDATE bannerlar SET sira=? WHERE id=?", (i, bid))
    conn.commit(); conn.close()
    return {"mesaj": "Sıra güncellendi"}

@app.delete("/api/bannerlar/{bid}")
def banner_sil(bid: int, payload=Depends(admin_gerek)):
    conn = get_db()
    row = conn.execute("SELECT resim_url FROM bannerlar WHERE id=?", (bid,)).fetchone()
    if row and row["resim_url"]:
        dosya = BASE_DIR / row["resim_url"].lstrip("/")
        if dosya.exists():
            try: dosya.unlink()
            except: pass
    conn.execute("DELETE FROM bannerlar WHERE id=?", (bid,))
    conn.commit(); conn.close()
    return {"mesaj": "Silindi"}

@app.post("/api/bannerlar/{bid}/resim")
async def banner_resim(bid: int, dosya: UploadFile = File(...),
                       payload=Depends(admin_gerek)):
    uzanti = Path(dosya.filename or "").suffix.lower()
    if uzanti not in (".jpg", ".jpeg", ".png", ".webp", ".gif"):
        raise HTTPException(400, "jpg/png/webp/gif kabul edilir")
    icerik = await dosya.read()
    if len(icerik) > 15 * 1024 * 1024:
        raise HTTPException(400, "Maks 15MB")

    ad = f"banner_{bid}_{uuid.uuid4().hex[:8]}{uzanti}"
    hedef = UPLOAD_DIR / ad
    
    # Pillow ile optimize et (GIF hariç)
    if PIL_VAR and uzanti != ".gif":
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(icerik))
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        # Banner boyutuna göre max genişlik
        max_w = 1920
        if img.width > max_w:
            oran = max_w / img.width
            img = img.resize((max_w, int(img.height * oran)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="WEBP", quality=88, optimize=True)
        hedef = UPLOAD_DIR / f"banner_{bid}_{uuid.uuid4().hex[:8]}.webp"
        with open(hedef, "wb") as f:
            f.write(buf.getvalue())
    else:
        with open(hedef, "wb") as f:
            f.write(icerik)

    url = f"/static/uploads/{hedef.name}"
    # Eski resmi sil
    eski = get_db().execute("SELECT resim_url FROM bannerlar WHERE id=?", (bid,)).fetchone()
    if eski and eski["resim_url"] and eski["resim_url"] != url:
        eski_d = BASE_DIR / eski["resim_url"].lstrip("/")
        if eski_d.exists():
            try: eski_d.unlink()
            except: pass

    conn = get_db()
    conn.execute("UPDATE bannerlar SET resim_url=? WHERE id=?", (url, bid))
    conn.commit(); conn.close()
    return {"url": url}

# ── SEO: robots.txt ─────────────────────────────────────────────────────────
@app.get("/robots.txt")
def robots():
    from fastapi.responses import PlainTextResponse
    return PlainTextResponse(
        "User-agent: *\n"
        "Allow: /\n"
        "Disallow: /api/\n"
        "Disallow: /static/uploads/\n\n"
        "Sitemap: https://portfoygayrimenkul.com.tr/sitemap.xml\n"
    )

# ── SEO: sitemap.xml ─────────────────────────────────────────────────────────
@app.get("/sitemap.xml")
def sitemap_xml():
    conn = get_db()
    portfoyler = conn.execute(
        "SELECT id, guncelleme FROM portfoyler WHERE durum='Aktif' ORDER BY guncelleme DESC"
    ).fetchall()
    bloglar = conn.execute(
        "SELECT slug, guncelleme FROM blog_yazilari WHERE durum='Yayında' ORDER BY guncelleme DESC"
    ).fetchall()
    conn.close()
    base = "https://portfoygayrimenkul.com.tr"
    urls = [
        f"  <url><loc>{base}/</loc><changefreq>daily</changefreq><priority>1.0</priority></url>"
    ]
    for p in portfoyler:
        urls.append(
            f"  <url><loc>{base}/?ilan={p['id']}</loc>"
            f"<lastmod>{str(p['guncelleme'])[:10]}</lastmod>"
            f"<priority>0.8</priority></url>"
        )
    for b in bloglar:
        urls.append(
            f"  <url><loc>{base}/?blog={b['slug']}</loc>"
            f"<lastmod>{str(b['guncelleme'])[:10]}</lastmod>"
            f"<priority>0.6</priority></url>"
        )
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n'
        + "\n".join(urls)
        + "\n</urlset>"
    )
    return Response(content=xml, media_type="application/xml")

# ── Blog Endpoint'leri ────────────────────────────────────────────────────────
@app.get("/api/blog")
def blog_listele(durum: str = "", payload=Depends(token_coz)):
    conn = get_db()
    is_admin = payload and payload.get("rol") == "admin"
    if is_admin and not durum:
        rows = conn.execute(
            "SELECT b.*, k.ad_soyad as yazar FROM blog_yazilari b "
            "LEFT JOIN kullanicilar k ON b.yazar_id=k.id ORDER BY b.olusturma DESC"
        ).fetchall()
    elif is_admin and durum:
        rows = conn.execute(
            "SELECT b.*, k.ad_soyad as yazar FROM blog_yazilari b "
            "LEFT JOIN kullanicilar k ON b.yazar_id=k.id WHERE b.durum=? ORDER BY b.olusturma DESC",
            (durum,)
        ).fetchall()
    else:
        rows = conn.execute(
            "SELECT b.*, k.ad_soyad as yazar FROM blog_yazilari b "
            "LEFT JOIN kullanicilar k ON b.yazar_id=k.id WHERE b.durum='Yayında' ORDER BY b.olusturma DESC"
        ).fetchall()
    conn.close()
    result = []
    for r in rows:
        d = dict(r)
        d["etiketler"] = json.loads(d.get("etiketler") or "[]")
        result.append(d)
    return result

@app.get("/api/blog/{slug}")
def blog_detay_api(slug: str):
    conn = get_db()
    row = conn.execute(
        "SELECT b.*, k.ad_soyad as yazar FROM blog_yazilari b "
        "LEFT JOIN kullanicilar k ON b.yazar_id=k.id WHERE b.slug=?", (slug,)
    ).fetchone()
    conn.close()
    if not row:
        raise HTTPException(404, "Yazı bulunamadı")
    d = dict(row)
    d["etiketler"] = json.loads(d.get("etiketler") or "[]")
    return d

@app.post("/api/blog")
def blog_ekle(b: BlogGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    slug = slug_olustur(b.baslik)
    mevcut = conn.execute("SELECT id FROM blog_yazilari WHERE slug=?", (slug,)).fetchone()
    if mevcut:
        slug = slug + "-" + uuid.uuid4().hex[:4]
    yazar = conn.execute("SELECT id FROM kullanicilar WHERE email=?", (payload["sub"],)).fetchone()
    conn.execute(
        "INSERT INTO blog_yazilari (baslik,slug,icerik,ozet,etiketler,kapak_resim,durum,yazar_id) "
        "VALUES (?,?,?,?,?,?,?,?)",
        (b.baslik, slug, b.icerik, b.ozet,
         json.dumps(b.etiketler, ensure_ascii=False),
         b.kapak_resim, b.durum,
         yazar["id"] if yazar else None)
    )
    bid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
    conn.commit()
    conn.close()
    return {"id": bid, "slug": slug, "mesaj": "Yazı oluşturuldu"}

@app.put("/api/blog/{bid}")
def blog_guncelle(bid: int, b: BlogGiren, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute(
        "UPDATE blog_yazilari SET baslik=?,icerik=?,ozet=?,etiketler=?,"
        "kapak_resim=?,durum=?,guncelleme=datetime('now') WHERE id=?",
        (b.baslik, b.icerik, b.ozet,
         json.dumps(b.etiketler, ensure_ascii=False),
         b.kapak_resim, b.durum, bid)
    )
    conn.commit()
    conn.close()
    return {"mesaj": "Yazı güncellendi"}

@app.delete("/api/blog/{bid}")
def blog_sil(bid: int, payload=Depends(admin_gerek)):
    conn = get_db()
    conn.execute("DELETE FROM blog_yazilari WHERE id=?", (bid,))
    conn.commit()
    conn.close()
    return {"mesaj": "Yazı silindi"}

@app.post("/api/blog/{bid}/kapak")
async def blog_kapak_yukle(bid: int, dosya: UploadFile = File(...), payload=Depends(admin_gerek)):
    from PIL import Image, ImageOps
    import io as _io

    uzanti = Path(dosya.filename or "").suffix.lower()
    if uzanti not in (".jpg", ".jpeg", ".png", ".webp"):
        raise HTTPException(400, "Sadece jpg/png/webp")

    icerik = await dosya.read()
    if len(icerik) > 12 * 1024 * 1024:
        raise HTTPException(400, "Dosya 12MB'dan küçük olmalı")

    try:
        img = Image.open(_io.BytesIO(icerik))
        img = ImageOps.exif_transpose(img)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        # Google og:image / Twitter Card için ideal: 1200x630 (1.91:1)
        hedef_w, hedef_h = 1200, 630
        hedef_oran = hedef_w / hedef_h
        kaynak_oran = img.width / img.height
        if kaynak_oran > hedef_oran:
            yeni_h = img.height
            yeni_w = int(yeni_h * hedef_oran)
            sol = (img.width - yeni_w) // 2
            img = img.crop((sol, 0, sol + yeni_w, yeni_h))
        else:
            yeni_w = img.width
            yeni_h = int(yeni_w / hedef_oran)
            ust = (img.height - yeni_h) // 2
            img = img.crop((0, ust, yeni_w, ust + yeni_h))
        if img.width > hedef_w:
            img = img.resize((hedef_w, hedef_h), Image.LANCZOS)

        ad = f"blog_{bid}_{uuid.uuid4().hex[:8]}.webp"
        hedef = UPLOAD_DIR / ad
        img.save(hedef, "WEBP", quality=87, method=4)
    except Exception as e:
        raise HTTPException(400, f"Resim işlenemedi: {str(e)[:150]}")

    url = f"/static/uploads/{ad}"
    conn = get_db()
    conn.execute("UPDATE blog_yazilari SET kapak_resim=? WHERE id=?", (url, bid))
    conn.commit()
    conn.close()
    return {"url": url}

# ── Blog İçerik Görseli — otomatik boyutlandırma, kare/dikdörtgen seçimi ──────
BLOG_BOYUT_ORANLARI = {
    "kare":        (1080, 1080),
    "dikdortgen":  (1200, 800),
    "genis":       (1600, 700),
    "orijinal":    None,
}

@app.post("/api/blog/icerik-resim")
async def blog_icerik_resim_yukle(
    dosya: UploadFile = File(...),
    boyut: str = "dikdortgen",
    konum: str = "ortali",
    payload=Depends(admin_gerek)
):
    """Blog içeriğine eklenecek görsel: otomatik boyutlandırma + konum bilgisiyle döner."""
    from PIL import Image, ImageOps
    import io as _io

    uzanti = Path(dosya.filename or "").suffix.lower()
    if uzanti not in (".jpg", ".jpeg", ".png", ".webp"):
        raise HTTPException(400, "Sadece jpg/png/webp kabul edilir")

    icerik = await dosya.read()
    if len(icerik) > 12 * 1024 * 1024:
        raise HTTPException(400, "Dosya 12MB'dan küçük olmalı")

    if konum not in ("basta", "ortali", "sonda"):
        konum = "ortali"
    if boyut not in BLOG_BOYUT_ORANLARI:
        boyut = "dikdortgen"

    try:
        img = Image.open(_io.BytesIO(icerik))
        img = ImageOps.exif_transpose(img)  # EXIF rotasyon düzelt
        if img.mode in ("RGBA", "P") and uzanti in (".jpg", ".jpeg"):
            img = img.convert("RGB")

        hedef_boyut = BLOG_BOYUT_ORANLARI[boyut]
        if hedef_boyut:
            # Cover-fit crop: oranı koru, taşanı kırp
            hedef_w, hedef_h = hedef_boyut
            hedef_oran = hedef_w / hedef_h
            kaynak_oran = img.width / img.height

            if kaynak_oran > hedef_oran:
                # Kaynak daha geniş — yüksekliği baz al, genişliği kırp
                yeni_h = img.height
                yeni_w = int(yeni_h * hedef_oran)
                sol = (img.width - yeni_w) // 2
                img = img.crop((sol, 0, sol + yeni_w, yeni_h))
            else:
                # Kaynak daha dar/uzun — genişliği baz al, yüksekliği kırp
                yeni_w = img.width
                yeni_h = int(yeni_w / hedef_oran)
                ust = (img.height - yeni_h) // 2
                img = img.crop((0, ust, yeni_w, ust + yeni_h))

            # Hedef boyuttan büyükse küçült (üst sınır)
            if img.width > hedef_w:
                img = img.resize((hedef_w, hedef_h), Image.LANCZOS)
        else:
            # Orijinal oran — sadece çok büyükse küçült (max 1600px genişlik)
            if img.width > 1600:
                oran = 1600 / img.width
                img = img.resize((1600, int(img.height * oran)), Image.LANCZOS)

        # Kaydet (webp = küçük boyut, kaliteli)
        kayit_uzanti = ".webp"
        ad = f"blog_icerik_{uuid.uuid4().hex[:10]}{kayit_uzanti}"
        hedef_yol = UPLOAD_DIR / ad
        img.save(hedef_yol, "WEBP", quality=85, method=4)

    except Exception as e:
        raise HTTPException(400, f"Resim işlenemedi: {str(e)[:150]}")

    url = f"/static/uploads/{ad}"
    return {
        "url": url,
        "boyut": boyut,
        "konum": konum,
        "genislik": img.width,
        "yukseklik": img.height,
    }

# ─── Static Dosyalar ──────────────────────────────────────────────────────────
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

@app.get("/{full_path:path}")
def spa_fallback(full_path: str):
    index = BASE_DIR / "static" / "index.html"
    if index.exists():
        return FileResponse(str(index))
    return JSONResponse({"mesaj": "API çalışıyor"}, status_code=200)

# ─── Başlat ────────────────────────────────────────────────────────────────────
init_db()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)

# ═══════════════════════════════════════════════════════════════════════════════
# FAZ 3 — SEO + BLOG + HARİTA
# ═══════════════════════════════════════════════════════════════════════════════
