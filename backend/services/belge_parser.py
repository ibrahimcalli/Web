"""Belge parse yardımcıları."""
from __future__ import annotations
import re
from typing import Any

def docx_parse(dosya_yolu: str) -> dict:
    """docx dosyasından portföy alanlarını çıkar."""
    try:
        from docx import Document
        doc = Document(dosya_yolu)
        tablo_verisi = {}
        metin_bloklari = []

        # Tabloları oku
        for tablo in doc.tables:
            for satir in tablo.rows:
                cells = [c.text.strip() for c in satir.cells]
                if len(cells) >= 2 and cells[0] and cells[1]:
                    tablo_verisi[cells[0].strip(":")] = cells[1]

        # Paragrafları oku
        for para in doc.paragraphs:
            metin = para.text.strip()
            if metin:
                metin_bloklari.append(metin)

        return _tablo_to_portfoy(tablo_verisi, metin_bloklari)
    except Exception as e:
        return {"hata": str(e)}

def html_parse(icerik: str) -> dict:
    """HTML içeriğinden portföy alanlarını çıkar."""
    from html.parser import HTMLParser
    import html as html_module

    tablo_verisi = {}
    metin_bloklari = []

    # Tablo hücrelerini yakala
    td_pattern = re.findall(r'<td[^>]*>(.*?)</td>', icerik, re.DOTALL | re.IGNORECASE)
    temiz = [re.sub(r'<[^>]+>', '', h).strip() for h in td_pattern]
    temiz = [html_module.unescape(t) for t in temiz if t]

    for i in range(0, len(temiz) - 1, 2):
        if temiz[i] and temiz[i+1]:
            tablo_verisi[temiz[i].strip(":")] = temiz[i+1]

    # h2 başlıklar (bölüm bilgisi)
    for m in re.finditer(r'<h[1-3][^>]*>(.*?)</h[1-3]>', icerik, re.DOTALL | re.IGNORECASE):
        t = re.sub(r'<[^>]+>', '', m.group(1)).strip()
        if t:
            metin_bloklari.append(t)

    return _tablo_to_portfoy(tablo_verisi, metin_bloklari)

def _tablo_to_portfoy(tablo: dict, metinler: list) -> dict:
    """Tablo verisi + metin → portföy dict."""
    def bul(*anahtarlar):
        for a in anahtarlar:
            for k, v in tablo.items():
                if a.lower() in k.lower():
                    return v.strip()
        return ""

    # Başlık: ilk metin satırı genellikle işletme adı + ilan tipi
    baslik = ""
    for m in metinler:
        if len(m) > 5 and not m.startswith("IC "):
            baslik = m
            break

    # Kategori tespiti
    kat_str  = bul("Kategori", "Kategori")
    alt_str  = bul("Ilan Tipi", "İlan Tipi", "Alt Kategori")
    durum_str = bul("Durum", "Statü")

    # Kategori → ana_kat / alt_kat çıkar
    ana_kat, alt_kat, ilan_tipi = "Konut", "Satılık", ""
    for k in KATEGORILER:
        if k.lower() in kat_str.lower():
            ana_kat = k
            break
    for a in KATEGORILER.get(ana_kat, []):
        if a.lower() in kat_str.lower() or a.lower() in alt_str.lower():
            alt_kat = a
            break

    # Fiyat: ilk büyük sayıyı yakala
    fiyat = ""
    for m in metinler:
        match = re.search(r'[\d.,]+\s*(TL|EUR|USD|\$|€)?', m)
        if match and any(c.isdigit() for c in match.group()):
            sayi = re.sub(r'[.,\s]', '', match.group())
            if len(sayi) >= 4:
                fiyat = match.group().strip()
                break

    # GPS
    gps = bul("GPS", "Konum", "google")
    gps_match = re.search(r'q=([\d.]+),([\d.]+)', gps)
    if gps_match:
        gps = f"{gps_match.group(1)},{gps_match.group(2)}"

    # Açıklama + saha notu
    aciklama = ""
    saha_notu = ""
    yakalanmis = False
    for m in metinler:
        if "açıklama" in m.lower() or "ACIKLAMA" in m:
            yakalanmis = True
            continue
        if "saha" in m.lower() or "SAHA" in m:
            yakalanmis = False
            saha_notu = m
            continue
        if yakalanmis and not aciklama:
            aciklama = m

    # Dinamik alanlar
    alanlar = {}
    alan_map = {
        "net_m2":        ["Net M2", "Net m2", "Net M²"],
        "brut_m2":       ["Brut M2", "Brüt m2", "Brüt M²"],
        "oda_sayisi":    ["Oda Sayisi", "Oda Sayısı"],
        "bina_kati":     ["Bina Kati", "Bina Katı"],
        "bulundugu_kat": ["Kat", "Bulunduğu Kat"],
        "bina_yasi":     ["Bina Yasi", "Bina Yaşı"],
        "banyo_sayisi":  ["Banyo Sayisi", "Banyo Sayısı"],
        "isitma":        ["Isitma", "Isıtma"],
        "tapu_durumu":   ["Tapu Durumu"],
        "cephe":         ["Cephe"],
        "esyali":        ["Esyali", "Eşyalı"],
        "balkon":        ["Balkon"],
        "asansor":       ["Asansor", "Asansör"],
        "otopark":       ["Otopark"],
        "site_icinde":   ["Site Icinde", "Site İçinde"],
        "kullanim":      ["Kullanim", "Kullanım"],
        "takas":         ["Takas"],
        "ada":           ["Ada"],
        "parsel":        ["Parsel"],
        "krediye_uygun": ["Krediye Uygun"],
        "kimden":        ["Kimden"],
        "ozellikler":    ["Deger Katan", "Değer Katan", "Özellikler"],
    }
    for hedef, kaynaklar in alan_map.items():
        deger = bul(*kaynaklar)
        if deger:
            alanlar[hedef] = deger

    musteri_ad   = bul("Ad Soyad", "Sahip", "Müşteri")
    musteri_tel  = bul("Telefon", "Tel")
    musteri_mail = bul("E-posta", "Mail", "Email")
    il_ilce      = bul("Il / Ilce", "İl / İlçe", "İl/İlçe")
    il, ilce     = "Muğla", "Fethiye"
    if "/" in il_ilce:
        parcalar = [p.strip() for p in il_ilce.split("/")]
        il  = parcalar[0] if parcalar[0] else il
        ilce = parcalar[1] if len(parcalar) > 1 and parcalar[1] else ilce
    mahalle = bul("Mahalle", "Bölge", "Semt")

    return {
        "baslik":       baslik,
        "ana_kategori": ana_kat,
        "alt_kategori": alt_kat,
        "ilan_tipi":    ilan_tipi,
        "fiyat":        fiyat,
        "il":           il,
        "ilce":         ilce,
        "mahalle":      mahalle,
        "gps":          gps,
        "aciklama":     aciklama,
        "saha_notu":    saha_notu,
        "musteri_ad":   musteri_ad,
        "musteri_tel":  musteri_tel,
        "musteri_mail": musteri_mail,
        "alanlar":      alanlar,
        "kaynak":       "belge_import",
    }

# ─── Pydantic Modeller ─────────────────────────────────────────────────────────